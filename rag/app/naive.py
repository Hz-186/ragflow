import logging
import re
import os
from functools import reduce
from io import BytesIO
from timeit import default_timer as timer
from typing import Any, Callable
from docx import Document
from docx.opc.pkgreader import _SerializedRelationships, _SerializedRelationship
from docx.table import Table as DocxTable
from docx.text.paragraph import Paragraph
from docx.opc.oxml import parse_xml
from markdown import markdown
from PIL import Image
from common.token_utils import num_tokens_from_string

from common.constants import LLMType, MAXIMUM_PAGE_NUMBER
from api.db.services.llm_service import LLMBundle
from api.db.joint_services.tenant_model_service import (
    ensure_mineru_from_env,
    ensure_opendataloader_from_env,
    ensure_paddleocr_from_env,
    get_composite_model_name_by_id,
    get_first_provider_model_name,
    resolve_model_config,
    get_tenant_default_model_by_type,
)
from rag.utils.file_utils import extract_embed_file, extract_links_from_pdf, extract_links_from_docx, extract_html
from deepdoc.parser import DocxParser, EpubParser, ExcelParser, HtmlParser, JsonParser, MarkdownElementExtractor, MarkdownParser, PdfParser, TxtParser
from deepdoc.parser.figure_parser import VisionFigureParser, vision_figure_parser_docx_wrapper_naive, vision_figure_parser_pdf_wrapper
from deepdoc.parser.pdf_parser import PlainParser, VisionParser
from deepdoc.parser.docling_parser import DoclingParser
from deepdoc.parser.tcadp_parser import TCADPParser
from common.float_utils import normalize_overlapped_percent
from common.parser_config_utils import has_mineru_options, normalize_layout_recognizer
from common.text_utils import normalize_arabic_presentation_forms
from rag.nlp import (
    concat_img,
    find_codec,
    naive_merge,
    naive_merge_with_images,
    naive_merge_docx,
    rag_tokenizer,
    tokenize_chunks,
    tokenize_chunks_with_positions,
    doc_tokenize_chunks_with_images,
    tokenize_table,
    append_context2table_image4pdf,
    tokenize_chunks_with_images,
)  # noqa: F401


def _is_short_header(text, max_tokens=50):
    """
    检查指定文本是否为较短的 Markdown 标题 —— 短标题识别器。
    识别出的短标题不会被单独切成孤立切片，下游会强制把它与下方正文合并。

    Args:
        text (str): 待检查的文本内容，如 "# 第一章 概述"。
        max_tokens (int, optional): 判断标题是否为“短标题”的最大 Token 数量上限。默认值为 50。
    Returns:
        bool: 如果文本是符合长度限制的 Markdown 标题则返回 True，否则返回 False。
    """
    # 1. 基础判空：如果文本为 None、空字符串或全空白字符，直接返回 False
    if not text or not text.strip():
        return False

    # 2. 正则匹配 Markdown 标题语法：去除前后空白后，匹配以 1 到 6 个 '#' 开头且后接至少一个空白字符
    if not re.match(r"^#{1,6}\s+", text.strip()):
        return False

    # 3. Token 计数检查：计算文本的 Token 数量，判断是否严格小于设定的最大阈值
    return num_tokens_from_string(text) < max_tokens


def _normalize_section_text_for_rtl_presentation_forms(sections):
    """
    规范化章节（sections）数据中的 RTL（从右到左）文本，统一阿拉伯语字形表达形式。
    把文本部分的阿拉伯语表达形式（Arabic Presentation Forms）字符做 NFKC 归一，
    避免影响后续分词、向量化与全文检索匹配；章节的容器结构（元组/列表/字符串
    及附加元数据）原样保留。各格式（docx/pdf/markdown/txt/html/epub/json）解析
    出的 sections 都会过一遍这个函数。

    Args:
        sections (list | tuple | None): 待规范化的章节列表或元组。每个章节元素可以是：
            - tuple: 如 (text, image, ...) 等包含文本及其他元数据的元组
            - list: 如 [text, ...] 等包含文本及其他元数据的列表
            - str: 纯文本字符串

    Returns:
        list | tuple | None: 文本部分已被规范化后的章节结构；若传入空数据则直接原样返回。
    """
    # 1. 基础判空：若传入的 sections 为空或 None，直接原样返回
    if not sections:
        return sections

    normalized_sections = []
    # 2. 遍历每个章节元素，根据其具体的数据容器类型进行针对性处理
    for section in sections:
        # 2.1 处理元组类型的章节（例如：(text, image_info)）
        if isinstance(section, tuple):
            if not section:
                normalized_sections.append(section)  # 空元组原样保留
                continue
            text = section[0]  # 提取元组首项的文本内容
            normalized_text = normalize_arabic_presentation_forms(text)  # 规范化阿拉伯语/RTL 文本
            normalized_sections.append((normalized_text, *section[1:]))  # 保持元组结构及后续附加数据不变
            continue

        # 2.2 处理列表类型的章节（例如：[text, ...]）
        if isinstance(section, list):
            if not section:
                normalized_sections.append(section)  # 空列表原样保留
                continue
            text = section[0]  # 提取列表首项的文本内容
            normalized_text = normalize_arabic_presentation_forms(text)  # 规范化阿拉伯语/RTL 文本
            normalized_sections.append([normalized_text, *section[1:]])  # 保持列表结构及后续附加数据不变
            continue

        # 2.3 处理纯字符串或其他标量类型的章节
        normalized_sections.append(normalize_arabic_presentation_forms(section))

    # 3. 返回规范化处理完成的所有章节列表
    return normalized_sections


def _merge_excel_items(items, chunk_token_num=128):
    """把同一工作表的连续行按 token 预算攒成切片 —— Excel 行合并器。

    Excel 解析器先把表格拆成「一行一条」的文本，本函数负责把同一工作表里
    相邻的行用换行符拼回较大的切片，避免一行产生一个细碎切片。

    输入数据的样子：
        items —— 「(行文本, 位置五元组)」列表（ExcelParser 默认模式的产物，
        每行文本已拼成「表头：值; 表头：值」形式；表头行本身不单列，
        行号从 2 起）：
            [
                ("姓名：张三; 年龄：25; 城市：北京", (0, 2, 2, 1, 3)),
                ("姓名：李四; 年龄：30; 城市：上海", (0, 3, 3, 1, 3)),
            ]
            位置五元组是 (工作表序号, 起始行, 结束行, 起始列, 结束列)，
            工作表序号从 0 数，行列号用于前端检索命中后回表高亮。
            另：工作表名不含 "sheet" 时，行尾还会追加 " ——表名" 后缀
        chunk_token_num —— 单切片 token 上限；<=0 表示「不合并，原样返回」
            （html4excel 模式下每块表格已是完整 HTML，就是走这条路）

    返回值的样子：
        仍是「(文本, 位置五元组)」列表，只是同表相邻行已用换行合并，
        位置的行/列范围扩到覆盖被合并的所有行：
            [
                ("姓名：张三; 年龄：25; 城市：北京\\n姓名：李四; 年龄：30; 城市：上海",(0, 2, 3, 1, 3)),
                ...,
            ]
    """
    if not items:
        return []
    if chunk_token_num <= 0:
        return items  # 不合并模式：原样返回

    merged = []
    cur_text = ""  # 正在攒的切片文本
    cur_pos = None  # 正在攒的切片位置（行/列范围随合并不断外扩）
    cur_tokens = 0  # 正在攒的切片 token 数

    for text, pos in items:
        sheet_idx, r1, r2, c1, c2 = pos
        tok = num_tokens_from_string(text)
        same_sheet = cur_pos is not None and cur_pos[0] == sheet_idx
        # 切片封箱条件：已在攒内容，且（换到了另一个工作表，或加进这一行会超预算）
        if cur_text and (not same_sheet or cur_tokens + tok > chunk_token_num):
            merged.append((cur_text, cur_pos))  # 当前切片封箱
            cur_text = ""
            cur_pos = None
            cur_tokens = 0

        # 新切片开张：当前没有任何在攒内容时，这一行自己成为切片的第一行
        if not cur_text:
            cur_text = text
            cur_pos = (sheet_idx, r1, r2, c1, c2)
            cur_tokens = tok
            continue

        # 并入在攒切片：文本用换行拼接，位置取两者的并集（最小起始、最大结束）
        cur_text = cur_text + "\n" + text
        cur_pos = (
            sheet_idx,
            min(cur_pos[1], r1),
            max(cur_pos[2], r2),
            min(cur_pos[3], c1),
            max(cur_pos[4], c2),
        )
        cur_tokens += tok

    if cur_text:
        merged.append((cur_text, cur_pos))  # 循环结束后把最后一片收尾
    return merged


def by_deepdoc(filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    """用自研 DeepDOC 版面识别模型解析 PDF —— DeepDOC 解析入口。

    输入参数：
        filename —— 文件名/磁盘路径；若传了 binary 则不读盘，直接用字节流
        binary —— 文件原始字节流，如 b"%PDF-1.7..."；从对象存储取回的文件都走这条路
        from_page / to_page —— 只解析这个页码区间（分片任务用，0 基）
        lang —— 文档语言（"Chinese"/"English"）；在本函数里只转给视觉图片
            增强包装器，DeepDOC 自身的表格识别会按页面内容自动探测中英文
        callback —— task_executor 给的进度回调 callback(progress, msg)；
            progress 取 0~1 表正常进度，-1 表失败
        pdf_cls —— 可替换的 PDF 解析器类；其他解析方法（paper 等）可传自己的子类，
            缺省用本文件的 Pdf
        **kwargs —— tenant_id、parser_config 等，原样转给视觉图片增强包装器

    返回值的样子（三元组，后续由 chunk() 统一消费）：
        sections —— 正文段落流：
            [("第一段正文", "@@1\\t12.0\\t583.0\\t100.0\\t200.0##"), ...]
            第二个元素是 @@页\\t左\\t右\\t上\\t下## 位置标签，供切片阶段裁页面截图
        tables —— 表格和图片列表：
            [((<PIL 截图>, "<table>...</table>" 或 [图注文字]), [(页, 左, 右, 上, 下)]), ...]
            （第二个元素是字符串=表格、是列表=图片，这是与下游的显式约定）
        pdf_parser —— 实际使用的 Pdf 解析器对象（后续按位置标签裁图用）
    """
    # 实例化解析器：外部传了替代类就用替代类，否则用本文件的 Pdf
    pdf_parser = pdf_cls() if pdf_cls else Pdf()
    # 跑完整流水线：OCR 渲染 → 版面分析 → 表格识别 → 文本合并，产出段落流与表格/图片
    sections, tables = pdf_parser(filename if binary is None else binary, from_page=from_page, to_page=to_page, callback=callback)

    # 图片增强：租户配置了视觉大模型时，给表格/图片里的插图生成语义描述并写回文本
    tables = vision_figure_parser_pdf_wrapper(
        tbls=tables,
        sections=sections,
        callback=callback,
        lang=lang,
        **kwargs,
    )
    return sections, tables, pdf_parser


def _dispatch_pdf_parser(parser_config: dict, opendataloader_llm_name=None, layout_recognize_override: str | None = None) -> tuple[Callable, str, Any, str, Any]:
    """根据解析配置挑选对应的 PDF 解析函数 —— PDF 解析器调度器。
    单独抽成函数是为了让「配置 → 解析器」的调度逻辑可以独立测试（issue #17114）。

    输入参数的样子：
        parser_config —— 文档的解析配置字典：
            {
                "layout_recognize": "DeepDOC",   # 也可能是 "PlainText"/"MinerU"，
                    # 甚至是前端存旧的 TenantModel UUID
                "mineru_lang": "English",         # mineru_* 开头的键是 MinerU 专属选项
                "chunk_token_num": 512,
            }
        opendataloader_llm_name —— 调用方已解析好的 OpenDataLoader 模型名；没有就传 None
        layout_recognize_override —— 已经解析好的 layout_recognize 值（比如经
            get_composite_model_name_by_id 把 UUID 转成的 "模型名@实例@提供商" 复合名）。
            传了就用它，避免重新去读 parser_config 里的旧 UUID 再归一化一遍

    返回值的样子（五元组）：
        (parser, name, layout_recognizer, opendataloader_llm_name, parser_model_name)
            parser —— 选中的解析函数（PARSERS 表里 by_deepdoc/by_mineru/... 之一；
                名字不认识且无 mineru_* 选项时兜底到 by_plaintext）
            name —— 归一化小写后的解析器名，如 "deepdoc"、"mineru"；
                若是未知名字则保持原小写串（此时 parser 已兜底为 by_plaintext，
                或带 mineru_* 选项时被改道为 by_mineru）
            layout_recognizer —— 归一化后的识别器名，如 "DeepDOC"、"plaintext"
            opendataloader_llm_name —— OpenDataLoader 模型名（入参没传且复合名里
                没提取到时保持原值）
            parser_model_name —— layout_recognize 以 @mineru/@paddleocr/
                @opendataloader/@somark/@mistral ocr 结尾时，为 "模型@实例@提供商"
                完整复合名原串（下游按三段拆分定位提供商实例），否则为 None
    """
    # 取 layout_recognize 原始值：优先用外部已解析好的，其次才读配置
    raw_layout_recognize = layout_recognize_override if layout_recognize_override is not None else parser_config.get("layout_recognize", "DeepDOC")
    # 归一化：复合名（如 "xxx@yyy@zzz@mineru"）按后缀识别出真正的解析器，
    # 并把完整复合名原样留在 parser_model_name 里供下游按三段拆分；普通名字原样返回
    layout_recognizer, parser_model_name = normalize_layout_recognizer(raw_layout_recognize)
    if layout_recognizer == "OpenDataLoader" and parser_model_name:
        opendataloader_llm_name = parser_model_name  # 复合名原串覆盖入参模型名

    # 把 "Plain Text" 归一成 "plaintext"，让下面的 PARSERS 查表能命中显式注册的
    # "plaintext" 条目，而不是落到 get 的兜底值——这很关键：下方 MinerU 兜底
    # 判断依赖「名字是否是已知关键字」，走兜底和走命中是两种语义
    if layout_recognizer == "Plain Text":
        layout_recognizer = "plaintext"
    name = layout_recognizer.strip().lower()
    parser = PARSERS.get(name, by_plaintext)

    # 修复 issue #17114：文档上存的 layout_recognize 是模型 id（如 TenantModel UUID）、
    # 匹配不到任何已知解析器名时，旧逻辑会落到 by_plaintext，把 id 当视觉模型去解析，
    # 报 ``Provider <empty> not found for model <id>`` 后崩溃。若配置里同时带着
    # mineru_* 专属选项，说明用户本意就是用 MinerU，那就改走 by_mineru 并打日志，
    # 而不是悄悄掩盖配置错误。
    # 保护条件：只有「名字不是已知关键字」时才兜底。像
    # {"layout_recognize": "Plain Text", "mineru_lang": "English"} 这种配置
    # 必须仍然尊重 PlainText，不能被偷偷改道去 MinerU。
    if name not in PARSERS and parser is by_plaintext and has_mineru_options(parser_config):
        logging.warning(
            "[naive] layout_recognize=%r does not match a known parser; falling back to MinerU because mineru_* options are set (see issue #17114).",
            layout_recognizer,
        )
        parser = by_mineru
        name = "mineru"

    return parser, name, layout_recognizer, opendataloader_llm_name, parser_model_name


def by_mineru(
    filename,
    binary=None,
    from_page=0,
    to_page=MAXIMUM_PAGE_NUMBER,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    mineru_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    """用租户配置的 MinerU OCR 模型解析 PDF —— MinerU 解析入口。

    输入参数的样子：
        filename / binary / from_page / to_page / lang / callback —— 同 by_deepdoc
        parse_method —— MinerU 的解析方式，chunk() 走 naive 调度时会改写成 "naive"
        mineru_llm_name —— 已解析好的模型名；"模型@实例@提供商" 复合名时这里
            传的是完整复合名（下游按三段拆分定位提供商实例），可为 None
        tenant_id —— 租户 id；没有它整个函数直接抛异常（找不到任何模型来源）
        **kwargs —— 透传给 parse_pdf，其中 vision_model 会被自动补上（见下）

    返回值的样子：
        (sections, tables, pdf_parser) 三元组，结构与 by_deepdoc 相同；
        模型未配置时抛 RuntimeError，解析中途失败时原样重抛底层异常，
        两条路都不会返回 None
    """
    pdf_parser = None
    if tenant_id:
        # 找模型名：入参没给时，先查租户模型表，查不到再用环境变量兜底注册
        if not mineru_llm_name:
            try:
                mineru_llm_name = get_first_provider_model_name(tenant_id, "MinerU", LLMType.OCR) or ensure_mineru_from_env(tenant_id)
            except Exception as e:  # 兜底失败只记日志，不中断
                logging.warning(f"fallback to env mineru: {e}")

        if mineru_llm_name:
            try:
                # 按模型名解析出完整配置并包装成 LLMBundle，mdl 就是真正的解析器实例
                ocr_model_config = resolve_model_config(tenant_id, LLMType.OCR, mineru_llm_name)
                ocr_model = LLMBundle(tenant_id=tenant_id, model_config=ocr_model_config, lang=lang)
                pdf_parser = ocr_model.mdl

                # 修复 issue #14869：租户配置了视觉大模型时，把它塞进 kwargs，
                # 让 MinerU 给图片切片补上 VLM 生成的语义描述（与 DeepDOC 的
                # VisionFigureParser 能力对齐）。尽力而为——没有视觉模型就静默跳过
                if "vision_model" not in kwargs:
                    try:
                        vision_model_config = get_tenant_default_model_by_type(tenant_id, LLMType.VISION)
                        kwargs["vision_model"] = LLMBundle(tenant_id=tenant_id, model_config=vision_model_config, lang=lang)
                    except Exception as vlm_err:
                        logging.info(f"[MinerU] no VISION model for tenant; skipping image VLM enhancement: {vlm_err}")

                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    lang=lang,
                    page_from=from_page,
                    page_to=min(to_page, MAXIMUM_PAGE_NUMBER),
                    **kwargs,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle MinerU ({mineru_llm_name}): {e}")
                raise

    raise RuntimeError("MinerU model not found or not configured.")


def by_docling(filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    """用 Docling 引擎解析 PDF —— Docling 解析入口。

    运行参数全部来自环境变量：
        DOCLING_OUTPUT_DIR     —— 中间产物输出目录（空串=引擎默认位置）
        DOCLING_DELETE_OUTPUT  —— 解析完是否删除中间产物（默认 1=删）
        DOCLING_SERVER_URL     —— 独立 Docling 服务地址（空串=本地跑）

    输入参数的样子：同 by_deepdoc（parse_method 从 kwargs 里取，默认 "raw"）。

    返回值的样子：
        (sections, tables, pdf_parser) 三元组，结构与 by_deepdoc 相同；
        Docling 未安装时通过 callback(-1, ...) 报错并返回 (None, None, parser)，
        调用方需自行判空。
    """
    pdf_parser = DoclingParser()
    parse_method = kwargs.get("parse_method", "raw")

    # 环境检查：Docling 依赖没装齐时直接报失败，不进入解析
    if not pdf_parser.check_installation():
        if callback:
            callback(-1, "Docling not found.")
        return None, None, pdf_parser

    sections, tables = pdf_parser.parse_pdf(
        filepath=filename,
        binary=binary,
        callback=callback,
        output_dir=os.environ.get("DOCLING_OUTPUT_DIR", ""),
        delete_output=bool(int(os.environ.get("DOCLING_DELETE_OUTPUT", 1))),
        docling_server_url=os.environ.get("DOCLING_SERVER_URL", ""),
        parse_method=parse_method,
    )
    return sections, tables, pdf_parser


def by_opendataloader(
    filename,
    binary=None,
    from_page=0,
    to_page=MAXIMUM_PAGE_NUMBER,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    opendataloader_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    """用租户配置的 OpenDataLoader OCR 模型解析 PDF —— OpenDataLoader 解析入口。

    与 by_mineru 的区别：解析失败不抛异常，而是报失败并返回 (None, None, None)。

    输入参数的样子：同 by_deepdoc，另多三个可选解析开关（从 kwargs 里挑出来透传）：
        hybrid —— 是否混合模式；image_output —— 是否输出图片；
        sanitize —— 是否清洗文本。没传的开关不会出现在调用里。

    返回值的样子：
        (sections, tables, pdf_parser) 三元组，结构与 by_deepdoc 相同；
        找不到模型或解析失败时返回 (None, None, None) 并 callback(-1, ...)。
    """
    if tenant_id:
        # 找模型名：入参没给时，先查租户模型表，查不到再用环境变量兜底注册
        if not opendataloader_llm_name:
            try:
                opendataloader_llm_name = get_first_provider_model_name(tenant_id, "OpenDataLoader", LLMType.OCR) or ensure_opendataloader_from_env(tenant_id)
            except Exception as e:  # 兜底失败只记日志，不中断
                logging.warning(f"fallback to env opendataloader: {e}")

        if opendataloader_llm_name:
            try:
                # 按模型名解析出完整配置并包装成 LLMBundle，mdl 就是真正的解析器实例
                ocr_model_config = resolve_model_config(tenant_id, LLMType.OCR, opendataloader_llm_name)
                ocr_model = LLMBundle(tenant_id=tenant_id, model_config=ocr_model_config, lang=lang)
                pdf_parser = ocr_model.mdl
                # 只挑 OpenDataLoader 认识的三个开关透传，其余 kwargs 不进解析器
                parse_options = {k: kwargs[k] for k in ("hybrid", "image_output", "sanitize") if k in kwargs}
                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    **parse_options,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle OpenDataLoader ({opendataloader_llm_name}): {e}")

    if callback:
        callback(-1, "OpenDataLoader not found.")
    return None, None, None


def by_tcadp(filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, lang="Chinese", callback=None, pdf_cls=None, **kwargs):
    """用腾讯云文档解析（TCADP）云端 API 解析 PDF —— TCADP 解析入口。

    中间产物目录由环境变量 TCADP_OUTPUT_DIR 控制（空串=引擎默认位置）。

    输入参数的样子：同 by_deepdoc（云解析不需要页码区间与语言参数，不读它们）。

    返回值的样子：
        (sections, tables, parser) 三元组，结构与 by_deepdoc 相同；
        云端 API 未配置时通过 callback(-1, ...) 报错并返回 (None, None, parser)。
    """
    tcadp_parser = TCADPParser()

    # 环境检查：腾讯云 API 配置缺失时直接报失败，不发起解析
    if not tcadp_parser.check_installation():
        callback(-1, "TCADP parser not available. Please check Tencent Cloud API configuration.")
        return None, None, tcadp_parser

    sections, tables = tcadp_parser.parse_pdf(filepath=filename, binary=binary, callback=callback, output_dir=os.environ.get("TCADP_OUTPUT_DIR", ""), file_type="PDF")
    return sections, tables, tcadp_parser


def by_paddleocr(
    filename,
    binary=None,
    from_page=0,
    to_page=MAXIMUM_PAGE_NUMBER,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    paddleocr_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    """用租户配置的 PaddleOCR 模型解析 PDF —— PaddleOCR 解析入口。

    输入参数的样子：同 by_deepdoc。**kwargs 原样透传给 parse_pdf。

    返回值的样子：
        (sections, tables, pdf_parser) 三元组，结构与 by_deepdoc 相同；
        失败路径一律返回 (None, None, None)，但报错方式分三种：
        有租户、解析中途失败 → 只记日志不报错；有租户、但查不到任何模型来源
        → 静默返回（查询本身抛异常除外，那会先打一条兜底 warning）；
        没有租户 → callback(-1, ...) 报错。
    """
    pdf_parser = None
    if tenant_id:
        # 找模型名：入参没给时，先查租户模型表，查不到再用环境变量兜底注册
        if not paddleocr_llm_name:
            try:
                paddleocr_llm_name = get_first_provider_model_name(tenant_id, "PaddleOCR", LLMType.OCR) or ensure_paddleocr_from_env(tenant_id)
            except Exception as e:  # 兜底失败只记日志，不中断
                logging.warning(f"fallback to env paddleocr: {e}")

        if paddleocr_llm_name:
            try:
                # 按模型名解析出完整配置并包装成 LLMBundle，mdl 就是真正的解析器实例
                ocr_model_config = resolve_model_config(tenant_id, LLMType.OCR, paddleocr_llm_name)
                ocr_model = LLMBundle(tenant_id=tenant_id, model_config=ocr_model_config, lang=lang)
                pdf_parser = ocr_model.mdl
                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    **kwargs,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle PaddleOCR ({paddleocr_llm_name}): {e}")

        return None, None, None

    if callback:
        callback(-1, "PaddleOCR not found.")
    return None, None, None


def by_somark(
    filename,
    binary=None,
    from_page=0,
    to_page=MAXIMUM_PAGE_NUMBER,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    somark_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    """用租户配置的 SoMark OCR 模型解析 PDF —— SoMark 解析入口。

    注意：模型名只有「入参 → 环境变量兜底」两步来源，
    不查租户模型表，这是它与 by_mineru/by_paddleocr 的区别。

    输入参数的样子：同 by_deepdoc。**kwargs 原样透传给 parse_pdf。

    返回值的样子：
        (sections, tables, pdf_parser) 三元组，结构与 by_deepdoc 相同；
        解析失败或找不到模型时 callback(-1, ...) 报错并返回 (None, None, None)。
    """
    pdf_parser = None
    if tenant_id:
        # 找模型名：入参没给时，用环境变量兜底注册（延迟导入避免循环依赖）
        if not somark_llm_name:
            try:
                from api.db.joint_services.tenant_model_service import ensure_somark_from_env

                somark_llm_name = ensure_somark_from_env(tenant_id)
            except Exception as e:
                logging.warning(f"fallback to env somark: {e}")

        if somark_llm_name:
            try:
                # 按模型名解析出完整配置并包装成 LLMBundle，mdl 就是真正的解析器实例
                ocr_model_config = resolve_model_config(tenant_id, LLMType.OCR, somark_llm_name)
                ocr_model = LLMBundle(tenant_id=tenant_id, model_config=ocr_model_config, lang=lang)
                pdf_parser = ocr_model.mdl
                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    **kwargs,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle SoMark ({somark_llm_name}): {e}")
                if callback:
                    callback(-1, f"Failed to parse pdf via SoMark ({somark_llm_name}): {e}")
                return None, None, None

    if callback:
        callback(-1, "SoMark not found.")
    return None, None, None


def by_mistral_ocr(
    filename,
    binary=None,
    from_page=0,
    to_page=MAXIMUM_PAGE_NUMBER,
    lang="Chinese",
    callback=None,
    pdf_cls=None,
    parse_method: str = "raw",
    mistral_ocr_llm_name: str | None = None,
    tenant_id: str | None = None,
    **kwargs,
):
    """用租户配置的 Mistral OCR 模型解析 PDF —— Mistral OCR 解析入口。

    模型名来源与 by_somark 相同（入参 → 环境变量兜底两步）。

    输入参数的样子：同 by_deepdoc；页码区间 from_page/to_page 会传给解析器
    （by_paddleocr/by_somark 等不传页码）。**kwargs 透传给 parse_pdf。

    返回值的样子：
        (sections, tables, pdf_parser) 三元组，结构与 by_deepdoc 相同；
        解析失败或找不到模型时 callback(-1, ...) 报错并返回 (None, None, None)。
    """
    pdf_parser = None
    if tenant_id:
        # 找模型名：入参没给时，用环境变量兜底注册（延迟导入避免循环依赖）
        if not mistral_ocr_llm_name:
            try:
                from api.db.joint_services.tenant_model_service import ensure_mistral_ocr_from_env

                mistral_ocr_llm_name = ensure_mistral_ocr_from_env(tenant_id)
            except Exception as e:
                logging.warning(f"fallback to env mistral ocr: {e}")

        if mistral_ocr_llm_name:
            try:
                # 按模型名解析出完整配置并包装成 LLMBundle，mdl 就是真正的解析器实例
                ocr_model_config = resolve_model_config(tenant_id, LLMType.OCR, mistral_ocr_llm_name)
                ocr_model = LLMBundle(tenant_id=tenant_id, model_config=ocr_model_config, lang=lang)
                pdf_parser = ocr_model.mdl
                # 尽力而为的图片描述：把租户的视觉大模型塞进 kwargs，让 Mistral OCR
                # 抽出的图片也能拿到 VLM 生成的语义描述（与 MinerU/DeepDOC 能力对齐）。
                # 租户没配视觉模型就静默跳过
                if "vision_model" not in kwargs:
                    try:
                        vision_model_config = get_tenant_default_model_by_type(tenant_id, LLMType.VISION)
                        kwargs["vision_model"] = LLMBundle(tenant_id=tenant_id, model_config=vision_model_config, lang=lang)
                    except Exception as vlm_err:
                        logging.info(f"[Mistral OCR] no vision model for tenant; skipping figure description: {vlm_err}")
                sections, tables = pdf_parser.parse_pdf(
                    filepath=filename,
                    binary=binary,
                    callback=callback,
                    parse_method=parse_method,
                    from_page=from_page,
                    to_page=to_page,
                    lang=lang,
                    **kwargs,
                )
                return sections, tables, pdf_parser
            except Exception as e:
                logging.error(f"Failed to parse pdf via LLMBundle Mistral OCR ({mistral_ocr_llm_name}): {e}")
                if callback:
                    callback(-1, f"Failed to parse pdf via Mistral OCR ({mistral_ocr_llm_name}): {e}")
                return None, None, None

    if callback:
        callback(-1, "Mistral OCR not found.")
    return None, None, None


def by_plaintext(filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, callback=None, **kwargs):
    """不做版面识别解析 PDF —— 纯文本解析入口（也是兜底解析器）：默认用
    PlainParser 直接抽文字流；

    kwargs["layout_recognizer"] 填了具体视觉模型名时，
    改用该视觉大模型逐页「看图转文字」（VisionParser）。

    输入参数的样子：
        filename / binary / from_page / to_page / callback —— 同 by_deepdoc
        kwargs 里额外读两个键：
            "layout_recognizer" —— 空/缺省/"plaintext" 走纯文本；
                填视觉模型名（如 "qwen-vl-max"）走视觉模式
            "tenant_id" —— 视觉模式必需（要按租户查模型配置），纯文本模式不需要

    返回值的样子：
        (sections, tables, pdf_parser) 三元组，结构与 by_deepdoc 相同；
        视觉模式缺 tenant_id 时抛 ValueError。
    """
    layout_recognizer = (kwargs.get("layout_recognizer") or "").strip()
    if (not layout_recognizer) or layout_recognizer.replace(" ", "").lower() == "plaintext":
        pdf_parser = PlainParser()  # 纯文本模式：直接抽文字流
    else:
        # 视觉模式：按名字解析出租户的视觉大模型，交给 VisionParser 逐页识图
        tenant_id = kwargs.get("tenant_id")
        if not tenant_id:
            raise ValueError("tenant_id is required when using vision layout recognizer")
        vision_model_config = resolve_model_config(tenant_id, LLMType.VISION, layout_recognizer)
        vision_model = LLMBundle(
            tenant_id,
            model_config=vision_model_config,
            lang=kwargs.get("lang", "Chinese"),
        )
        pdf_parser = VisionParser(vision_model=vision_model, **kwargs)

    sections, tables = pdf_parser(filename if binary is None else binary, from_page=from_page, to_page=to_page, callback=callback)
    return sections, tables, pdf_parser


# PDF 解析器注册表：键是 layout_recognize 归一化小写后的名字，值是解析函数。
# 查不到名字时由 _dispatch_pdf_parser 兜底到 by_plaintext
PARSERS = {
    "deepdoc": by_deepdoc,
    "mineru": by_mineru,
    "docling": by_docling,
    "opendataloader": by_opendataloader,
    "tcadp parser": by_tcadp,
    "paddleocr": by_paddleocr,
    "somark": by_somark,
    "mistral ocr": by_mistral_ocr,
    "plaintext": by_plaintext,  # 默认兜底解析器
}


class Docx(DocxParser):
    """docx 解析器：把 Word 文档读成「文本/图片/表格」三混段落流 —— docx 解析主力。

    由 chunk() 的 docx 分支调用：Docx()(filename, binary) 产出段落流后，
    交给 naive_merge_docx 做合并切片。表格会转成 HTML 并带上「文档名 > 章节标题」
    的位置面包屑，图片保留 PIL 原图。
    """

    def __init__(self):
        pass

    def __clean(self, line):
        """清洗一行文本 —— 全角空格换成半角空格，再去掉首尾空白。"""
        line = re.sub(r"\u3000", " ", line).strip()
        return line

    def __get_nearest_title(self, table_index, filename):
        """找出某张表格所处的章节标题路径 —— 表格面包屑生成器。

        从表格位置往回倒着找最近的标题段落（Heading 1~7 样式），再一路向上
        找齐所有父级标题，拼成「文档名 > 一级标题 > 二级标题」的面包屑字符串，
        写进表格 HTML 的 <caption>，让脱离版式的表格切片仍知道自己属于哪一章。

        输入参数的样子：
            table_index —— 目标是文档里的第几张表（0 基，按文档顺序数）
            filename —— 文件名（含扩展名），去掉扩展名后作为面包屑的第一段，
                如 "季度报告.docx" → "季度报告"

        返回值的样子：
            "季度报告 > 第三章 销售数据 > 3.1 区域明细" —— 用 " > " 连接的标题路径；
            表格前找不到任何标题时返回空字符串 ""
        """
        import re
        from docx.text.paragraph import Paragraph

        titles = []
        blocks = []

        # 文档名 = 去掉扩展名的文件名；去完是空串就用兜底名
        doc_name = re.sub(r"\.[a-zA-Z]+$", "", filename)
        if not doc_name:
            doc_name = "Untitled Document"

        # 按文档顺序收集所有块：段落记 ("p", 序号, 段落对象)，表格只记占位
        try:
            for i, block in enumerate(self.doc._element.body):
                if block.tag.endswith("p"):  # 段落块
                    p = Paragraph(block, self.doc)
                    blocks.append(("p", i, p))
                elif block.tag.endswith("tbl"):  # 表格块（对象稍后按需再取）
                    blocks.append(("t", i, None))
        except Exception as e:
            logging.error(f"Error collecting blocks: {e}")
            return ""

        # 定位目标表格在 blocks 里的序号：按顺序数表格，数到第 table_index 张为止
        target_table_pos = -1
        table_count = 0
        for i, (block_type, pos, _) in enumerate(blocks):
            if block_type == "t":
                if table_count == table_index:
                    target_table_pos = pos
                    break
                table_count += 1

        if target_table_pos == -1:
            return ""  # 目标表格不存在，返回空面包屑

        # 从文档末尾往前倒着扫，找表格之前最近的一个标题段落
        nearest_title = None
        for i in range(len(blocks) - 1, -1, -1):
            block_type, pos, block = blocks[i]
            if pos >= target_table_pos:  # 跳过表格之后的块
                continue

            if block_type != "p":
                continue

            # 样式名形如 "Heading 2" 就是标题段落（大小写不敏感）
            if block.style and block.style.name and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                try:
                    level_match = re.search(r"(\d+)", block.style.name)
                    if level_match:
                        level = int(level_match.group(1))
                        if level <= 7:  # 只认 1~7 级标题
                            title_text = block.text.strip()
                            if title_text:  # 空标题不算数
                                nearest_title = (level, title_text)
                                break
                except Exception as e:
                    logging.error(f"Error parsing heading level: {e}")

        if nearest_title:
            titles.append(nearest_title)  # 最近的标题先入列
            current_level = nearest_title[0]

            # 继续向上找父级标题：只要当前层级大于 1，就再往前找一个层级更小的
            # （允许跨级，比如 4 级标题的父级可以直接是 1 级）
            while current_level > 1:
                found = False
                for i in range(len(blocks) - 1, -1, -1):
                    block_type, pos, block = blocks[i]
                    if pos >= target_table_pos:  # 跳过表格之后的块
                        continue

                    if block_type != "p":
                        continue

                    if block.style and re.search(r"Heading\s*(\d+)", block.style.name, re.I):
                        try:
                            level_match = re.search(r"(\d+)", block.style.name)
                            if level_match:
                                level = int(level_match.group(1))
                                # 层级数字更小 = 更高的标题（Heading 1 最高）
                                if level < current_level:
                                    title_text = block.text.strip()
                                    if title_text:  # 空标题不算数
                                        titles.append((level, title_text))
                                        current_level = level
                                        found = True
                                        break
                        except Exception as e:
                            logging.error(f"Error parsing parent heading: {e}")

                if not found:  # 找不到更高级的标题就收手
                    break

            # 标题按层级从小到大排（1 级在前），再拼上文档名得到面包屑
            titles.sort(key=lambda x: x[0])
            hierarchy = [doc_name] + [t[1] for t in titles]
            return " > ".join(hierarchy)

        return ""

    def __call__(self, filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER):
        """把 docx 按文档顺序读成「文本/图片/表格」三混段落流 —— 主解析流程。

        输入参数的样子：
            filename —— 磁盘路径；传了 binary 就直接用字节流，不读盘
            binary —— docx 原始字节流，如 b"PK\\x03\\x04..."（docx 本质是 zip）
            from_page / to_page —— 页码区间（0 基，右开）

        返回值的样子：
            [(文本, 图片, 表格), ...] 三元组列表，按文档顺序排列：
                [("第一章 概述...", None, None),           # 纯文本段落
                 ("", <PIL.Image 图片>, None),              # 图片块
                 ("", None, "<table><caption>...</table>")] # 表格块（HTML）
            每行三者只占其一，其余两项是 None/空串
        """
        self.doc = Document(filename) if binary is None else Document(BytesIO(binary))
        pn = 0  # 当前页码（靠文档里的分页标记累加）
        lines = []  # 段落流结果（先存 dict，最后统一转三元组）
        last_image = None  # 暂存的「无文字段落里的图片」，等下一个内容块来时再落盘
        table_idx = 0  # 当前是第几张表（给面包屑定位用）

        def flush_last_image():
            """把暂存的图片落成一个独立图片块（新内容块到来前调用，保证顺序）。"""
            nonlocal last_image, lines
            if last_image is not None:
                lines.append({"text": "", "image": last_image, "table": None, "style": "Image"})
                last_image = None

        # 按文档顺序遍历正文里的每一个块（段落或表格）
        for block in self.doc._element.body:
            if pn > to_page:
                break  # 已越过目标页区间，提前收工

            if block.tag.endswith("p"):  # —— 段落块 ——
                p = Paragraph(block, self.doc)

                if from_page <= pn < to_page:  # 只处理目标页区间内的段落
                    text = p.text.strip()
                    style_name = p.style.name if p.style else ""

                    if text:
                        if style_name == "Caption":
                            # 图注/表注段落：吸附紧挨在它前面的图片，
                            # 让「图 + 图注」合成同一行，切片时不分离
                            former_image = None

                            if lines and lines[-1].get("image") and lines[-1].get("style") != "Caption":
                                # 前一行就是图片块 → 弹出它，图片挂到本行。
                                # 注：style != "Caption" 这个条件实际恒为真——
                                # Caption 行入列时不带 style 键，全文件只有
                                # flush_last_image 会写 style="Image"
                                former_image = lines[-1].get("image")
                                lines.pop()

                            elif last_image is not None:
                                # 前一个块是「暂存的无文字图片」→ 取过来挂到本行
                                former_image = last_image
                                last_image = None

                            lines.append(
                                {
                                    "text": self.__clean(text),
                                    "image": former_image if former_image else None,
                                    "table": None,
                                }
                            )

                        else:
                            # 普通文本段落：先把暂存图片落盘（保持先后顺序），再加文本行
                            flush_last_image()
                            lines.append(
                                {
                                    "text": self.__clean(text),
                                    "image": None,
                                    "table": None,
                                }
                            )

                            # 段落里内嵌了图片（文字与图同段）→ 追加一个独立图片块
                            current_image = self.get_picture(self.doc, p)
                            if current_image is not None:
                                lines.append(
                                    {
                                        "text": "",
                                        "image": current_image,
                                        "table": None,
                                    }
                                )

                    else:
                        # 无文字段落：里面可能有图 → 先暂存，等下一个内容块来时
                        # 由 flush_last_image 落盘（给「图注紧跟图片」的吸附留机会）
                        current_image = self.get_picture(self.doc, p)
                        if current_image is not None:
                            last_image = current_image

                # 页码统计：扫描段落里每个 run 的分页标记
                for run in p.runs:
                    xml = run._element.xml
                    if "lastRenderedPageBreak" in xml:  # Word 渲染时留下的自动分页
                        pn += 1
                        continue
                    if "w:br" in xml and 'type="page"' in xml:  # 手动插入的分页符
                        pn += 1

            elif block.tag.endswith("tbl"):  # —— 表格块 ——
                if pn < from_page or pn > to_page:
                    table_idx += 1  # 区间外的表格不解析，但序号照常累加
                    continue

                flush_last_image()  # 表格前的暂存图片先落盘
                tb = DocxTable(block, self.doc)
                title = self.__get_nearest_title(table_idx, filename)  # 章节面包屑
                html = "<table>"
                if title:
                    html += f"<caption>Table Location: {title}</caption>"
                # 逐行转 HTML；相邻文字相同的单元格合并成 colspan，还原跨列单元格
                for r in tb.rows:
                    html += "<tr>"
                    col_idx = 0
                    try:
                        while col_idx < len(r.cells):
                            span = 1
                            c = r.cells[col_idx]
                            for j in range(col_idx + 1, len(r.cells)):
                                if c.text == r.cells[j].text:  # 相邻同文 = 跨列单元格
                                    span += 1
                                    col_idx = j
                                else:
                                    break
                            col_idx += 1
                            html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                    except Exception as e:
                        logging.warning(f"Error parsing table, ignore: {e}")
                    html += "</tr>"
                html += "</table>"
                lines.append({"text": "", "image": None, "table": html})
                table_idx += 1

        flush_last_image()  # 文档末尾若还有暂存图片，收尾落盘
        # dict 段落流转成 (文本, 图片, 表格) 三元组列表 —— 交给 naive_merge_docx
        new_line = [(line.get("text"), line.get("image"), line.get("table")) for line in lines]

        return new_line

    def to_markdown(self, filename=None, binary=None, inline_images: bool = True):
        """把 docx 整体转成 Markdown 文本 —— 走 mammoth + markdownify 两步转换。

        本函数使用了 mammoth 库（BSD 2-Clause 许可证）。
        与 __call__ 的分工：__call__ 产出三混段落流给切片流水线用；
        本方法产出一整篇 Markdown 文本，给需要全文视图的场景用。

        输入参数的样子：
            filename / binary —— 二选一，传了 binary 优先用字节流
            inline_images —— True 时图片转成 base64 data URL 内嵌在 Markdown 里；
                False 时图片直接丢弃

        返回值的样子：
            "# 标题\\n\\n正文段落...\\n![img_a1b2c3d4](data:image/png;base64,...)"
            （一个完整的 Markdown 字符串）
        """

        import base64
        import uuid

        import mammoth
        from markdownify import markdownify

        docx_file = BytesIO(binary) if binary is not None else open(filename, "rb")

        def _convert_image_to_base64(image):
            """mammoth 的图片回调：把文档内嵌图片转成 base64 data URL。"""
            try:
                with image.open() as image_file:
                    image_bytes = image_file.read()
                encoded = base64.b64encode(image_bytes).decode("utf-8")
                base64_url = f"data:{image.content_type};base64,{encoded}"

                alt_name = "image"
                alt_name = f"img_{uuid.uuid4().hex[:8]}"  # 随机 8 位后缀当替代文本

                return {"src": base64_url, "alt": alt_name}
            except Exception as e:
                logging.warning(f"Failed to convert image to base64: {e}")
                return {"src": "", "alt": "image"}

        try:
            # 第一步：mammoth 把 docx 转成 HTML（可选带内嵌图片）
            if inline_images:
                result = mammoth.convert_to_html(docx_file, convert_image=mammoth.images.img_element(_convert_image_to_base64))
            else:
                result = mammoth.convert_to_html(docx_file)

            html = result.value

            # 第二步：markdownify 把 HTML 转成 Markdown
            markdown_text = markdownify(html)
            return markdown_text

        finally:
            if binary is None:
                docx_file.close()  # 磁盘打开的文件句柄确保关闭


class Pdf(PdfParser):
    """DeepDOC 路线的 PDF 解析器：跑完「OCR → 版面 → 表格 → 文本合并」四步流水线。"""

    def __init__(self):
        super().__init__()

    def __call__(self, filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, zoomin=3, callback=None, separate_tables_figures=False):
        """跑完整解析流水线，产出段落流 + 表格/图片 —— PDF 解析主流程。

        输入参数的样子：
            filename / binary —— 二选一，传了 binary 优先用字节流
            from_page / to_page —— 页码区间（0 基，右开）
            zoomin —— 渲染放大倍数（越大识别越细，越慢），默认 3
            callback —— 进度回调 callback(progress, msg)，本流程把进度打到 0.63~0.67
            separate_tables_figures —— True 时表格和图片分开返回（三元组），
                False 时混在一个列表里（二元组）

        返回值的样子：
            separate_tables_figures=False（默认）→ 二元组 (sections, tbls)：
                sections —— [(段落文本, "@@页\\t左\\t右\\t上\\t下##"), ...]，
                    位置标签由 _line_tag 生成，供切片阶段裁页面截图
                tbls —— [((<PIL 截图>, "<table>HTML</table>" 或 [图注]), [(页, 左, 右, 上, 下)]), ...]
            separate_tables_figures=True → 三元组 (sections, tbls, figures)，
                tbls 只剩表格、figures 只剩图片，各自带位置
        """
        start = timer()
        first_start = start
        callback(msg="OCR started")
        # 第一步：把 PDF 页面渲染成图片并做 OCR 文字识别
        self.__images__(filename if binary is None else binary, zoomin, from_page, to_page, callback)
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.info("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))

        start = timer()
        # 第二步：版面分析，识别出文本/表格/图片/标题等区块
        self._layouts_rec(zoomin)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        # 第三步：表格结构识别，把表格区块还原成行列结构
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        start = timer()
        # 第四步：文本行合并，把零散文字块按阅读顺序拼成段落
        self._text_merge(zoomin=zoomin)
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))

        if separate_tables_figures:
            # 表格/图片分开抽出（各自带位置），正文做向下拼接收尾
            tbls, figures = self._extract_table_figure(True, zoomin, True, True, True)
            self._concat_downward()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls, figures
        else:
            # 表格/图片混装抽出，正文额外做一轮纵向合并再向下拼接
            tbls = self._extract_table_figure(True, zoomin, True, True)
            self._naive_vertical_merge()
            self._concat_downward()
            # self._final_reading_order_merge()
            # self._filter_forpages()
            logging.info("layouts cost: {}s".format(timer() - first_start))
            return [(b["text"], self._line_tag(b, zoomin)) for b in self.boxes], tbls


# 抓取 markdown 文档引用的远程图片时，最多跟随的 HTTP 重定向跳数
# （每一跳都要重新做 SSRF 安全校验，见 Markdown.load_images_from_urls）
MAX_IMAGE_REDIRECTS = 5


class Markdown(MarkdownParser):
    """markdown 解析器：按标题切节、抽表格、下载内嵌图片 —— markdown 解析主力。

    由 chunk() 的 markdown 分支调用，产出 (小节文本列表, 表格列表, 小节图片列表)。
    """

    def md_to_html(self, sections):
        """把一段 markdown 文本转成 BeautifulSoup 对象 —— 供超链接提取用。

        输入参数的样子：
            sections —— 小节文本（字符串），或「(文本, ...)」元组（取第一项文本），
                如 "## 参考\\n详见 [官网](https://example.com)"

        返回值的样子：
            BeautifulSoup 对象（可 find_all("a") 遍历链接）；
            输入为空或不是文本形态时返回空列表 []
        """
        if not sections:
            return []
        if isinstance(sections, type("")):
            text = sections
        elif isinstance(sections[0], type("")):
            text = sections[0]  # 元组形态：取第一项文本
        else:
            return []

        from bs4 import BeautifulSoup

        html_content = markdown(text)  # markdown → HTML
        soup = BeautifulSoup(html_content, "html.parser")
        return soup

    def get_hyperlink_urls(self, soup):
        """从 BeautifulSoup 对象里抽出所有超链接地址（去重集合）。"""
        if soup:
            return set([a.get("href") for a in soup.find_all("a") if a.get("href")])
        return []

    def extract_image_urls_with_lines(self, text):
        """抽出 markdown 文本里所有图片引用的地址及其所在行号 —— 图片引用收集器。

        输入参数的样子：
            text —— 去掉表格后的 markdown 正文（字符串）

        返回值的样子：
            [{"url": "img/logo.png", "line": 3}, {"url": "https://x.cn/a.png", "line": 10}, ...]
            line 是 0 基行号；同一 (地址, 行号) 只记一次
        """
        md_img_re = re.compile(r"!\[[^\]]*\]\(([^)\s]+)")  # ![说明](地址)
        # 遗留缺陷提醒：这个字符类写成了 \\s（反斜杠+字母 s 两个字符），
        # 实际排除了字母 s——含 s 的地址会在第一个 s 处被截断
        # （如 "https://..." 只截到 "http"）。截断的垃圾条目会留在结果里，
        # 完整地址实际靠下面第 ③ 步的 BeautifulSoup 全文解析补全
        html_img_re = re.compile(r'src=["\\\']([^"\\\'>\\s]+)', re.IGNORECASE)  # src="地址"
        urls = []
        seen = set()  # (地址, 行号) 去重表
        lines = text.splitlines()
        # ①② 逐行扫两种图片语法（② 的正则有截断缺陷，见上方提醒）
        for idx, line in enumerate(lines):
            for url in md_img_re.findall(line):
                if (url, idx) not in seen:
                    urls.append({"url": url, "line": idx})
                    seen.add((url, idx))
            for url in html_img_re.findall(line):
                if (url, idx) not in seen:
                    urls.append({"url": url, "line": idx})
                    seen.add((url, idx))

        # ③ 全文补捞：BeautifulSoup 既能解析属性换行的 <img> 标签，
        # 也能把 ② 里被截断的地址补成完整地址（靠 (地址, 行号) 去重共存）
        try:
            from bs4 import BeautifulSoup

            soup = BeautifulSoup(text, "html.parser")
            # 预先记下全文每个换行符的偏移量，用于把字符位置换算成行号
            newline_offsets = [m.start() for m in re.finditer(r"\n", text)] + [len(text)]
            for img_tag in soup.find_all("img"):
                src = img_tag.get("src")
                if not src:
                    continue

                # 把标签还原成字符串回原文里找位置，再换算成行号
                tag_str = str(img_tag)
                pos = text.find(tag_str)
                if pos == -1:
                    # 找不到完整标签串时，退而求其次用 src 地址本身定位
                    pos = max(text.find(src), 0)
                line_no = 0
                for i, off in enumerate(newline_offsets):
                    if pos <= off:
                        line_no = i
                        break
                if (src, line_no) not in seen:
                    urls.append({"url": src, "line": line_no})
                    seen.add((src, line_no))
        except Exception as e:
            logging.error("Failed to extract image urls: {}".format(e))
            pass

        return urls

    def load_images_from_urls(self, urls, cache=None):
        """按地址列表把图片下载/读取成 PIL 图片 —— 图片加载器（带缓存与防穿透）。

        输入参数的样子：
            urls —— 图片地址列表，远程和本地路径可混装：
                ["https://x.cn/a.png", "img/logo.png", ...]
            cache —— 跨小节共用的缓存 {地址: PIL 图或 None}；首次调用可不传，
                失败也会缓存成 None，同一地址不会重复下载

        返回值的样子：
            (images, cache) 二元组：
                images —— [<PIL.Image RGB 图>, ...]（只含成功加载的，按 urls 顺序）
                cache —— 更新后的缓存字典，调用方应接住下次传回来
        """
        import requests
        from pathlib import Path
        from urllib.parse import urljoin

        from common.ssrf_guard import assert_url_is_safe, pin_dns

        cache = cache or {}
        images = []
        for url in urls:
            if url in cache:  # 命中缓存：成功过的直接复用
                if cache[url]:
                    images.append(cache[url])
                continue
            img_obj = None
            try:
                if url.startswith(("http://", "https://")):
                    # SSRF 防线：图片地址来自用户上传的（不可信）文档，连接前必须
                    # 校验并锁定 DNS。否则 ![x](http://169.254.169.254/...) 这样的
                    # 引用会让服务器去请求内网服务/云元数据接口。
                    # 重定向手动逐跳跟随，每一跳都重新校验，
                    # 做法与 common/data_source/rss_connector.py 一致。
                    current_hostname, current_ip = assert_url_is_safe(url)
                    current_url = url
                    response = None
                    try:
                        for _ in range(MAX_IMAGE_REDIRECTS + 1):
                            # 开下一跳前先释放上一跳：stream=True 时连接要等正文
                            # 读完或 response 关闭才会归还连接池
                            if response is not None:
                                response.close()
                            # pin_dns：把请求钉死在校验过的 IP 上，
                            # 防止 DNS 重新解析被劫持到内网地址
                            with pin_dns(current_hostname, current_ip):
                                response = requests.get(current_url, stream=True, timeout=30, allow_redirects=False)
                            if response.status_code not in (301, 302, 303, 307, 308):
                                break  # 非重定向响应 → 到达最终地址
                            location = response.headers.get("Location")
                            if not location:
                                break  # 说是重定向却没给地址 → 放弃跟随
                            current_url = urljoin(current_url, location)
                            current_hostname, current_ip = assert_url_is_safe(current_url)  # 新跳重新校验
                        else:
                            raise ValueError(f"Exceeded {MAX_IMAGE_REDIRECTS} redirects fetching {url!r}")
                        # 只有 200 且内容类型确实是图片才收下
                        if response.status_code == 200 and response.headers.get("Content-Type", "").startswith("image/"):
                            img_obj = Image.open(BytesIO(response.content)).convert("RGB")
                    finally:
                        # 任何路径（含非图片、重定向超限）都要释放流式 response，
                        # 否则正文没读完的连接不会归还连接池
                        if response is not None:
                            response.close()
                else:
                    # 本地相对路径：文件存在才读
                    local_path = Path(url)
                    if local_path.exists():
                        img_obj = Image.open(url).convert("RGB")
                    else:
                        logging.warning(f"Local image file not found: {url}")
            except Exception as e:
                logging.error(f"Failed to download/open image from {url}: {e}")
            cache[url] = img_obj  # 成功失败都记缓存，避免重复尝试
            if img_obj:
                images.append(img_obj)
        return images, cache

    def __call__(self, filename, binary=None, separate_tables=True, delimiter=None, return_section_images=False):
        """把 markdown 解析成「按标题切好的小节文本流 + 表格 + 小节图片」—— 主解析流程。

        输入参数的样子：
            filename / binary —— 二选一，传了 binary 优先用字节流（自动探测编码）
            separate_tables —— True 时表格单独抽出成表项；
                chunk() 的 markdown 分支传 False（表格留在正文里一起切）
            delimiter —— 自定义切节分隔符正则，缺省按标题层级切节
            return_section_images —— True 时额外返回每个小节的图片列表（三元组）

        返回值的样子：
            return_section_images=True（chunk() 走这条）→ 三元组：
                sections —— [("小节正文", ""), ("## 第二章...", ""), ...]
                tbls —— [((None, "<table>...</table>"), ""), ...]
                    （表格先转成 HTML；第一项里的 None 表示无截图）
                section_images —— [<PIL 图或 None>, ...] 与 sections 下标一一对应
                    （小节内多张图已用 concat_img 纵向拼成一张）
            return_section_images=False → 只返回 (sections, tbls)
        """
        if binary is not None:
            encoding = find_codec(binary)  # 自动探测字节流编码
            txt = binary.decode(encoding, errors="ignore")
        else:
            with open(filename, "r") as f:
                txt = f.read()

        # 先把表格整体抽走（若配置了单独成表），剩下的正文再按标题切节
        remainder, tables = self.extract_tables_and_remainder(f"{txt}\n", separate_tables=separate_tables)
        parsing_text = remainder
        extractor = MarkdownElementExtractor(parsing_text)
        image_refs = self.extract_image_urls_with_lines(parsing_text)  # 全文图片引用（带行号）
        element_sections = extractor.extract_elements(delimiter, include_meta=True)

        sections = []
        section_images = []
        image_cache = {}  # 跨小节共用的图片缓存，同一地址只下载一次
        # 逐个小节：按行号范围圈出本小节引用的图片，下载后拼成一张图
        for element in element_sections:
            content = element["content"]
            start_line = element["start_line"]
            end_line = element["end_line"]
            urls_in_section = [ref["url"] for ref in image_refs if start_line <= ref["line"] <= end_line]
            imgs = []
            if urls_in_section:
                imgs, image_cache = self.load_images_from_urls(urls_in_section, image_cache)
            combined_image = None
            if imgs:
                combined_image = reduce(concat_img, imgs) if len(imgs) > 1 else imgs[0]
            sections.append((content, ""))
            section_images.append(combined_image)

        # 抽走的表格逐个转成 HTML 表项（形态对齐其他解析器的表格产物）
        tbls = []
        if separate_tables:
            for table in tables:
                tbls.append(((None, markdown(table, extensions=["markdown.extensions.tables"])), ""))
        if return_section_images:
            return sections, tbls, section_images
        return sections, tbls


def load_from_xml_v2(baseURI, rels_item_xml):
    """python-docx 关系加载函数的替换版：跳过指向 "NULL" 的坏关系 —— 兼容性补丁。
    修复原生版遇到 "../NULL"、"NULL" 这类无效目标时报
    "There is no item named 'word/NULL' in the archive" 的问题；
    在 chunk() 的 docx 分支被猴子补丁到 _SerializedRelationships.load_from_xml。

    输入参数的样子：
        baseURI —— 关系所在部件的基础 URI，如 "word/document.xml"
        rels_item_xml —— rels 关系文件的 XML 字节串，None 时返回空集合

    返回值的样子：
        _SerializedRelationships 实例（内部 _srels 是过滤后的关系列表）
    """
    srels = _SerializedRelationships()
    if rels_item_xml is not None:
        rels_elm = parse_xml(rels_item_xml)
        for rel_elm in rels_elm.Relationship_lst:
            # 跳过指向 "NULL" 的坏关系和 # 开头的文档内锚点
            if rel_elm.target_ref in ("../NULL", "NULL") or rel_elm.target_ref.startswith("#"):
                continue
            srels._srels.append(_SerializedRelationship(baseURI, rel_elm))
    return srels


def chunk(filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, lang="Chinese", callback=None, **kwargs):
    """通用（naive）切片总入口：按扩展名选解析器，把文件变成可检索的 ES 切片列表。

    支持 docx、pdf、excel(csv/xlsx)、txt/代码、markdown、html、epub、json、doc，
    由 task_executor.build_chunks 调用（parser_id="naive"/"general"）。

    输入参数的样子：
        filename —— 文件名（含扩展名），扩展名决定走哪个解析分支，如 "报告.docx"
        binary —— 文件原始字节流（task_executor 从对象存储取回后传入）
        from_page / to_page —— 页码区间（主要服务 PDF 分片任务，0 基）
        lang —— "Chinese"/"English"，影响分词策略
        callback —— 进度回调 callback(progress, msg)；-1 表示失败
        **kwargs 里常用的键：
            parser_config —— 切片配置，如 {
                "chunk_token_num": 512, "delimiter": "\\n!?。；！？",
                "children_delimiter": "；", "layout_recognize": "DeepDOC",
                "table_context_size": 0, "image_context_size": 0,
                "analyze_hyperlink": True, "overlapped_percent": 0,
                "html4excel": False, "hyperlink_urls": False,
            }
            tenant_id —— 租户 id（查视觉/OCR 模型用）
            kb_id —— 知识库 id
            is_root —— 递归防爆闸，默认 True；抓链接/挖内嵌文件的递归子调用传 False

    返回值的样子：
        ES 文档（切片）列表：
            [
                {
                    "docnm_kwd": "报告.docx",
                    "content_with_weight": "切片原文（表格/图片切片带上下文）",
                    "content_ltks": "粗粒度分词结果",
                    "content_sm_ltks": "细粒度分词结果",
                    "position_int": [...],          # 定位坐标
                    "image": <PIL 图>,              # 图片/表格/带图切片才有
                    ...
                },
                ...
            ]
        递归抓取的超链接内容和内嵌文件切片也混在同一列表里返回
    """
    urls = set()
    url_res = []

    lang = lang or "Chinese"
    is_english = lang.lower() == "english"  # 是否英文（影响下游分词器的选择）
    parser_config = kwargs.get("parser_config", {"chunk_token_num": 512, "delimiter": "\n!?。；！？", "layout_recognize": "DeepDOC", "analyze_hyperlink": True})

    # ===== 父子切片：解析 children_delimiter（子分隔符）=====
    # ① 编码绕一圈：前端传来的 "\n" 是「反斜杠+n」两个字符，
    #    经 unicode_escape 解码成真正的换行符；中间用 latin1 是为了逐字节保真
    child_deli = (parser_config.get("children_delimiter") or "").encode("utf-8").decode("unicode_escape").encode("latin1").decode("utf-8")
    cust_child_deli = re.findall(r"`([^`]+)`", child_deli)  # ② 抽出所有反引号包裹的「多字符子分隔符」
    child_deli = "|".join(re.sub(r"`([^`]+)`", "", child_deli))  # ③ 剩下的裸字符：每个字符各自是一个子分隔符，用 | 连成正则
    if cust_child_deli:
        cust_child_deli = sorted(set(cust_child_deli), key=lambda x: -len(x))  # ④ 去重 + 按长度降序（长的优先命中，与主分隔符规则一致）
        cust_child_deli = "|".join(re.escape(t) for t in cust_child_deli if t)
        child_deli += cust_child_deli  # ⑤ 拼成完整交替正则，随切片传给 split_with_pattern 做二次切分

    is_markdown = False
    table_context_size = max(0, int(parser_config.get("table_context_size", 0) or 0))
    image_context_size = max(0, int(parser_config.get("image_context_size", 0) or 0))

    doc = {"docnm_kwd": filename, "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))}
    doc["title_sm_tks"] = rag_tokenizer.fine_grained_tokenize(doc["title_tks"])
    res = []
    pdf_parser = None
    section_images = None

    # ===== 递归切片：内嵌文件只在「根调用」时挖一层 =====
    # is_root 是递归防爆闸：最外层调用默认 True，递归子调用显式传 False，
    # 保证「文件里内嵌的文件」只被挖一层，不会无限递归下去
    is_root = kwargs.get("is_root", True)
    embed_res = []
    if is_root:
        # 只有根调用才挖内嵌文件（防止嵌套递归）
        embeds = []
        if binary is not None:
            embeds = extract_embed_file(binary)  # 从 zip 结构（word/embeddings/ 等）和 OLE 结构（Ole10Native）里挖出内嵌的 docx/xlsx/图片等，按内容哈希去重
        else:
            raise Exception("Embedding extraction from file path is not supported.")

        # 对每个内嵌文件递归调用 chunk()，把它们的切片并入主文档结果
        for embed_filename, embed_bytes in embeds:
            try:
                # 注意 is_root=False：内嵌文件自己不再挖它的内嵌层
                sub_res = chunk(embed_filename, binary=embed_bytes, lang=lang, callback=callback, is_root=False, **kwargs) or []
                embed_res.extend(sub_res)
            except Exception as e:
                # 单个内嵌文件解析失败不影响主文档，记日志后继续
                error_msg = f"Failed to chunk embed {embed_filename}: {e}"
                logging.error(error_msg)
                if callback:
                    callback(0.05, error_msg)
                continue

    if re.search(r"\.docx$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        # ===== 递归切片：docx 超链接也只在根调用时抓一层 =====
        if parser_config.get("analyze_hyperlink", False) and is_root:
            urls = extract_links_from_docx(binary)  # 从 docx 的 rels 里抽出所有外部超链接
            for index, url in enumerate(urls):
                html_bytes, metadata = extract_html(url)  # 抓取链接内容（转成 HTML 字节）
                if not html_bytes:
                    continue  # 抓不到的链接直接跳过
                try:
                    # 优先用链接本身当文件名（按扩展名走对应解析器）
                    sub_url_res = chunk(url, html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
                except Exception as e:
                    # 链接文件名不被任何解析器认识时，退回按 .html 解析
                    logging.info(f"Failed to chunk url in registered file type {url}: {e}")
                    sub_url_res = chunk(f"{index}.html", html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
                url_res.extend(sub_url_res)  # 链接内容切片并入主文档结果

        # 猴子补丁：替换 python-docx 的关系加载函数，跳过指向 "NULL" 的坏关系，
        # 修复 "There is no item named 'word/NULL' in the archive" 报错
        # （见 https://github.com/python-openxml/python-docx/issues/1105#issuecomment-1298075246）
        _SerializedRelationships.load_from_xml = load_from_xml_v2

        # 解析成 (文本, 图片, 表格) 三混段落流，再做阿拉伯语字形归一化
        sections = Docx()(filename, binary)
        sections = _normalize_section_text_for_rtl_presentation_forms(sections)

        # 返回两样东西：
        #   chunks —— dict 切片列表（{"text", "image", "ck_type", ...}）
        #   images —— 图片切片在 chunks 里的下标列表（供视觉增强按图定位）
        chunks, images = naive_merge_docx(sections, int(parser_config.get("chunk_token_num", 128)), parser_config.get("delimiter", "\n!?。；！？"), table_context_size, image_context_size)

        # 图片增强：租户配了视觉大模型时，给图片切片补上 VLM 语义描述
        vision_figure_parser_docx_wrapper_naive(
            chunks=chunks,
            idx_lst=images,
            callback=callback,
            lang=lang,
            **kwargs,
        )

        callback(0.8, "Finish parsing.")
        st = timer()

        res.extend(doc_tokenize_chunks_with_images(chunks, doc, is_english, child_delimiters_pattern=child_deli, language=lang))
        logging.info("naive_merge({}): {}".format(filename, timer() - st))
        res.extend(embed_res)
        res.extend(url_res)
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        layout_recognize_raw = parser_config.get("layout_recognize", "DeepDOC")
        tenant_id = kwargs.get("tenant_id")
        if tenant_id and isinstance(layout_recognize_raw, str):
            try:
                layout_recognize_raw = get_composite_model_name_by_id(layout_recognize_raw)
            except LookupError:
                pass
        # 注意：这里不要再自己调 normalize_layout_recognizer()——
        # 调度器内部会对 layout_recognize_override 参数做归一化，
        # 这里重复调一遍会把解析出的复合名丢掉，弄断
        # @mineru/@paddleocr/@opendataloader/@somark/@mistral ocr 复合名的模型接线
        # （CodeRabbit 对 #17114 的第 5 条评审意见）
        opendataloader_llm_name = kwargs.pop("opendataloader_llm_name", None)
        # 把已解析好的 layout_recognize 传进去（get_composite_model_name_by_id
        # 已把 TenantModel UUID 转成 "模型@实例@提供商" 复合名），
        # 调度器才能从中提取出正确的 parser_model_name
        parser, name, layout_recognizer, opendataloader_llm_name, parser_model_name = _dispatch_pdf_parser(
            parser_config,
            opendataloader_llm_name,
            layout_recognize_override=layout_recognize_raw,
        )

        # ===== 递归切片：PDF 超链接同样只在根调用时抓一层 =====
        # 这里只是先收集链接，真正抓取与切片在本函数末尾统一进行
        if parser_config.get("analyze_hyperlink", False) and is_root:
            urls = extract_links_from_pdf(binary)  # 从 PDF 注释里抽出所有外部链接
        callback(0.1, "Start to parse.")
        if name == "mineru":
            kwargs["parse_method"] = "naive"

        sections, tables, pdf_parser = parser(
            filename=filename,
            binary=binary,
            from_page=from_page,
            to_page=to_page,
            lang=lang,
            callback=callback,
            layout_recognizer=layout_recognizer,
            mineru_llm_name=parser_model_name,
            paddleocr_llm_name=parser_model_name,
            opendataloader_llm_name=opendataloader_llm_name,
            somark_llm_name=parser_model_name,
            mistral_ocr_llm_name=parser_model_name,
            **kwargs,
        )
        sections = _normalize_section_text_for_rtl_presentation_forms(sections)

        if not sections and not tables:
            return []

        if table_context_size or image_context_size:
            tables = append_context2table_image4pdf(
                sections,
                tables,
                image_context_size,
                section_page_offset=from_page if name == "mineru" else 0,
            )

        if name in ["tcadp", "docling", "mineru", "paddleocr", "opendataloader", "somark", "mistral ocr"]:
            if int(parser_config.get("chunk_token_num", 0)) <= 0:
                parser_config["chunk_token_num"] = 0

        res = tokenize_table(tables, doc, is_english, language=lang)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(csv|xlsx?)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")

        # 表格文件也支持换解析器：先看是否选了腾讯云 TCADP
        layout_recognizer = parser_config.get("layout_recognize", "DeepDOC")
        if layout_recognizer == "TCADP Parser":
            table_result_type = parser_config.get("table_result_type", "1")
            markdown_image_response_type = parser_config.get("markdown_image_response_type", "1")
            tcadp_parser = TCADPParser(table_result_type=table_result_type, markdown_image_response_type=markdown_image_response_type)
            if not tcadp_parser.check_installation():
                callback(-1, "TCADP parser not available. Please check Tencent Cloud API configuration.")
                return res

            # 按扩展名告诉云端 API 这是 XLSX 还是 CSV
            file_type = "XLSX" if re.search(r"\.xlsx?$", filename, re.IGNORECASE) else "CSV"

            sections, tables = tcadp_parser.parse_pdf(filepath=filename, binary=binary, callback=callback, output_dir=os.environ.get("TCADP_OUTPUT_DIR", ""), file_type=file_type)
            sections = _normalize_section_text_for_rtl_presentation_forms(sections)
            parser_config["chunk_token_num"] = 0
            res = tokenize_table(tables, doc, is_english, language=lang)
            sections = []
            callback(0.8, "Finish parsing.")
        else:
            # 默认路线：本地 ExcelParser 解析
            excel_parser = ExcelParser()
            if parser_config.get("html4excel"):
                # html4excel 模式：表格按每 12 行切成 HTML 片段，
                # token 上限置 0 → 后续不再合并，每块 HTML 独立成切片
                excel_items = [item for item in excel_parser.html(binary, 12) if item and item[0]]
                parser_config["chunk_token_num"] = 0
            else:
                # 普通模式：逐行读成 (行文本, 位置五元组)，交给 _merge_excel_items 攒片
                excel_items = [item for item in excel_parser(binary) if item and item[0]]
            excel_items = [(normalize_arabic_presentation_forms(text), pos) for text, pos in excel_items]
            res.extend(
                tokenize_chunks_with_positions(
                    _merge_excel_items(excel_items, int(parser_config.get("chunk_token_num", 128))),
                    doc,
                    is_english,
                    child_delimiters_pattern=child_deli,
                    language=lang,
                )
            )
            sections = []
            callback(0.8, "Finish parsing.")

    elif re.search(r"\.(txt|py|js|java|c|cpp|h|php|go|ts|sh|cs|kt|sql)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        sections = TxtParser()(filename, binary, parser_config.get("chunk_token_num", 128), parser_config.get("delimiter", "\n!?;。；！？"))
        sections = _normalize_section_text_for_rtl_presentation_forms(sections)
        logging.info("TxtParser produced %d sections for %s", len(sections), filename)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(md|markdown|mdx)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        markdown_parser = Markdown(int(parser_config.get("chunk_token_num", 128)))
        sections, tables, section_images = markdown_parser(
            filename,
            binary,
            separate_tables=False,
            delimiter=parser_config.get("delimiter", "\n!?;。；！？"),
            return_section_images=True,
        )
        sections = _normalize_section_text_for_rtl_presentation_forms(sections)

        is_markdown = True

        try:
            vision_model_config = get_tenant_default_model_by_type(kwargs["tenant_id"], LLMType.VISION)
            vision_model = LLMBundle(kwargs["tenant_id"], vision_model_config, lang=lang)
            callback(0.2, "Visual model detected. Attempting to enhance figure extraction...")
        except Exception as e:
            logging.warning(f"Failed to detect figure extraction: {e}")
            vision_model = None

        if vision_model:
            # 逐小节处理图片：带图的小节交给视觉大模型生成语义描述，追加到正文
            for idx, (section_text, _) in enumerate(sections):
                images = []
                if section_images and len(section_images) > idx and section_images[idx] is not None:
                    images.append(section_images[idx])

                if images and len(images) > 0:
                    # 多张图先用 concat_img 纵向拼成一张再送视觉模型
                    combined_image = reduce(concat_img, images) if len(images) > 1 else images[0]
                    if section_images:
                        section_images[idx] = combined_image
                    else:
                        section_images = [None] * len(sections)
                        section_images[idx] = combined_image
                    markdown_vision_parser = VisionFigureParser(
                        vision_model=vision_model,
                        figures_data=[((combined_image, ["markdown image"]), [(0, 0, 0, 0, 0)])],
                        lang=lang,
                        **kwargs,
                    )
                    boosted_figures = markdown_vision_parser(callback=callback)
                    sections[idx] = (section_text + "\n\n" + "\n\n".join([fig[0][1] for fig in boosted_figures]), sections[idx][1])

        else:
            logging.warning("No visual model detected. Skipping figure parsing enhancement.")

        if parser_config.get("hyperlink_urls", False) and is_root:
            for idx, (section_text, _) in enumerate(sections):
                soup = markdown_parser.md_to_html(section_text)
                hyperlink_urls = markdown_parser.get_hyperlink_urls(soup)
                urls.update(hyperlink_urls)
        res = tokenize_table(tables, doc, is_english, language=lang)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(htm|html)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = HtmlParser()(filename, binary, chunk_token_num)
        sections = [(_, "") for _ in sections if _]
        sections = _normalize_section_text_for_rtl_presentation_forms(sections)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.epub$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = EpubParser()(filename, binary, chunk_token_num)
        sections = [(_, "") for _ in sections if _]
        sections = _normalize_section_text_for_rtl_presentation_forms(sections)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.(json|jsonl|ldjson)$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")
        chunk_token_num = int(parser_config.get("chunk_token_num", 128))
        sections = JsonParser(chunk_token_num)(binary)
        sections = [(_, "") for _ in sections if _]
        sections = _normalize_section_text_for_rtl_presentation_forms(sections)
        callback(0.8, "Finish parsing.")

    elif re.search(r"\.doc$", filename, re.IGNORECASE):
        callback(0.1, "Start to parse.")

        try:
            from tika import parser as tika_parser
        except Exception as e:
            callback(0.8, f"tika not available: {e}. Unsupported .doc parsing.")
            logging.warning(f"tika not available: {e}. Unsupported .doc parsing for {filename}.")
            return []

        binary = BytesIO(binary)
        doc_parsed = tika_parser.from_buffer(binary)
        if doc_parsed.get("content", None) is not None:
            sections = doc_parsed["content"].split("\n")
            sections = [(_, "") for _ in sections if _]
            sections = _normalize_section_text_for_rtl_presentation_forms(sections)
            callback(0.8, "Finish parsing.")
        else:
            error_msg = f"tika.parser got empty content from {filename}."
            callback(0.8, error_msg)
            logging.warning(error_msg)
            return []
    else:
        raise NotImplementedError("file type not supported yet(pdf, xlsx, doc, docx, txt supported)")

    st = timer()
    overlapped_percent = normalize_overlapped_percent(parser_config.get("overlapped_percent", 0))

    # ===== markdown 分支：按标题节合并切片（图文同步）=====
    # markdown 解析器已按标题把文档切成 section，这里把相邻小节攒到
    # chunk_token_num 以内，并把各小节携带的图片用 concat_img 纵向拼接
    if is_markdown:
        merged_chunks = []
        merged_images = []
        chunk_limit = max(0, int(parser_config.get("chunk_token_num", 128)))  # 单切片 token 上限

        current_text = ""  # 正在攒的切片文本
        current_tokens = 0  # 正在攒的切片 token 数
        current_image = None  # 正在攒的切片图片（多张小节图纵向拼接）

        for idx, sec in enumerate(sections):
            text = sec[0] if isinstance(sec, tuple) else sec
            sec_tokens = num_tokens_from_string(text)
            sec_image = section_images[idx] if section_images and idx < len(section_images) else None

            # 短标题不单独成片：当前攒着的若是短标题，强制与下一小节合并，
            # 避免产生「只有一个标题」的孤零零切片
            if current_text and not _is_short_header(current_text) and current_tokens + sec_tokens > chunk_limit:
                merged_chunks.append(current_text)  # 装满 → 当前切片封片
                merged_images.append(current_image)
                overlap_part = ""
                if overlapped_percent > 0:
                    # 重叠切片：把上一片尾部若干字符抄进下一片开头，避免跨片语义被切断
                    overlap_len = int(len(current_text) * overlapped_percent / 100)
                    if overlap_len > 0:
                        overlap_part = current_text[-overlap_len:]
                current_text = overlap_part
                current_tokens = num_tokens_from_string(current_text)
                current_image = current_image if overlap_part else None  # 有重叠文本才保留旧图

            if current_text:
                current_text += "\n" + text  # 小节并入正在攒的切片
            else:
                current_text = text
            current_tokens += sec_tokens

            if sec_image:
                # 小节带图 → 与已攒图片纵向拼接，保持图文顺序一致
                current_image = concat_img(current_image, sec_image) if current_image else sec_image

        if current_text:
            merged_chunks.append(current_text)  # 最后一片收尾
            merged_images.append(current_image)

        chunks = merged_chunks
        has_images = merged_images and any(img is not None for img in merged_images)

        if has_images:
            # 带图走「图文配对」包装流水线
            res.extend(tokenize_chunks_with_images(chunks, doc, is_english, merged_images, child_delimiters_pattern=child_deli, language=lang))
        else:
            # 无图走主干包装流水线（配置了子分隔符时同样触发父子切片）
            res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser, child_delimiters_pattern=child_deli, language=lang))
    else:
        # ===== 非 markdown 分支：按分隔符切段再合并 =====
        if section_images:
            if all(image is None for image in section_images):
                section_images = None  # 所有小节都没图 → 按无图路径处理

        if section_images:
            # 带图路径：段落流 + 图片列表一起进带图合并切片机（合并时图片纵向拼接）
            chunks, images = naive_merge_with_images(sections, section_images, int(parser_config.get("chunk_token_num", 128)), parser_config.get("delimiter", "\n!?。；！？"), overlapped_percent)
            res.extend(tokenize_chunks_with_images(chunks, doc, is_english, images, child_delimiters_pattern=child_deli, language=lang))
        else:
            # 无图路径：通用合并切片机（配置了反引号分隔符时走自定义规则模式）
            chunks = naive_merge(sections, int(parser_config.get("chunk_token_num", 128)), parser_config.get("delimiter", "\n!?。；！？"), overlapped_percent)

            res.extend(tokenize_chunks(chunks, doc, is_english, pdf_parser, child_delimiters_pattern=child_deli, language=lang))

    # ===== 递归切片收尾：把前面收集的超链接逐个抓取并切片 =====
    # 注意：docx 分支在自己的分支里已经抓取并 return 了，走不到这里；
    # 这段收尾实际服务 PDF / markdown 两条路径（只有这两个分支会收集 urls）。
    # 同样受 is_root 保护：链接内容递归进来时 is_root=False，不会再抓它的链接
    if urls and parser_config.get("analyze_hyperlink", False) and is_root:
        for index, url in enumerate(urls):
            html_bytes, metadata = extract_html(url)  # 抓取链接内容（转成 HTML 字节）
            if not html_bytes:
                continue  # 抓不到的链接直接跳过
            try:
                # 优先用链接本身当文件名（按扩展名走对应解析器）
                sub_url_res = chunk(url, html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
            except Exception as e:
                # 链接文件名不被任何解析器认识时，退回按 .html 解析
                logging.info(f"Failed to chunk url in registered file type {url}: {e}")
                sub_url_res = chunk(f"{index}.html", html_bytes, callback=callback, lang=lang, is_root=False, **kwargs)
            url_res.extend(sub_url_res)  # 链接内容切片并入主文档结果

    logging.info("naive_merge({}): {}".format(filename, timer() - st))

    if embed_res:
        res.extend(embed_res)
    if url_res:
        res.extend(url_res)
    # if table_context_size or image_context_size:
    #    attach_media_context(res, table_context_size, image_context_size)

    # 把 PDF 目录（大纲）作为临时元数据挂在第一个切片上；
    # task_executor 入库前会把它取出并持久化为文档元数据
    if res and pdf_parser and getattr(pdf_parser, "outlines", None):
        res[0]["__outline__"] = [{"title": title, "depth": depth} for title, depth, *_ in pdf_parser.outlines]

    return res


# 命令行手工调试入口：python -m rag.app.naive 文件路径（只解析前 10 页）
if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
