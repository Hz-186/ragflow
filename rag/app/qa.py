import logging
import re
import csv
from copy import deepcopy
from io import BytesIO
from timeit import default_timer as timer
from openpyxl import load_workbook

from common.constants import MAXIMUM_PAGE_NUMBER
from deepdoc.parser.utils import get_text
from rag.nlp import is_english, random_choices, qbullets_category, add_positions, has_qbullet, docx_question_level
from rag.nlp import rag_tokenizer, tokenize_table, concat_img
from deepdoc.parser import PdfParser, ExcelParser, DocxParser
from docx import Document
from markdown import markdown

from common.float_utils import get_float


class Excel(ExcelParser):
    """Excel 版 Q&A 解析器 —— 把「两列问答表」读成 (问题, 答案) 对。

    约定（与 chunk() 的 docstring 一致）：
        表格无表头，第 1 列是问题、第 2 列是答案，多 sheet 也可以。
        每个 sheet 的每一行取「前两个非空单元格」为一对 (q, a)，
        单元格里的 0 / 0.0 / False 是真实内容（不是空），不能当空值跳过。

    输入数据的样子：
        fnm / binary —— xlsx 文件路径或二进制内容

    输出数据的样子：
        res = [("什么是RAG？", "RAG是检索增强生成..."), ("如何部署？", "用docker..."), ...]
        同时把 self.is_english 置为 True/False（按抽样问题判断文档语言）。
    """

    def __call__(self, fnm, binary=None, callback=None):
        # 打开工作簿：给了路径走本地文件，给了 bytes 走内存流
        if not binary:
            wb = load_workbook(fnm)
        else:
            wb = load_workbook(BytesIO(binary))
        total = 0
        # 先统计所有 sheet 的总行数，作为回调进度条的满分
        for sheetname in wb.sheetnames:
            total += len(list(wb[sheetname].rows))

        res, fails = [], []
        for sheetname in wb.sheetnames:
            ws = wb[sheetname]
            rows = list(ws.rows)
            for i, r in enumerate(rows):
                q, a = "", ""
                for cell in r:
                    # 空格子跳过；0、0.0、False 是真实内容，不能跳过
                    if cell.value is None or str(cell.value).strip() == "":
                        continue
                    # 第一个非空单元格当问题 q，第二个当答案 a，第三个起不管
                    if not q:
                        q = str(cell.value)
                    elif not a:
                        a = str(cell.value)
                    else:
                        break
                # 问题答案都齐了才收进 res，缺一个记入失败行号（供回调展示）
                if q and a:
                    res.append((q, a))
                else:
                    fails.append(str(i + 1))
                # 每收集 999 对回调一次进度（前 60% 归解析，后面 40% 归下游任务）
                if len(res) % 999 == 0:
                    callback(len(res) * 0.6 / total, ("Extract pairs: {}".format(len(res)) + (f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        # 解析完再回调一次 60%，报告最终抽取对数与失败行号
        callback(0.6, ("Extract pairs: {}. ".format(len(res)) + (f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))
        # 去掉"问题/答案/Q/A"等前缀后，随机抽 30 个问题判断整份文档是英文还是中文，
        # 供 beAdoc 决定用 "Question:/Answer:" 还是 "问题：/回答：" 前缀
        self.is_english = is_english([rmPrefix(q) for q, _ in random_choices(res, k=30) if len(q) > 1])
        return res


class Pdf(PdfParser):
    """PDF 版 Q&A 解析器 —— 把「试卷/问答册式 PDF」读成 (问题, 答案, 截图, 坐标) 列表。

    处理链（一句话）：OCR 出全部文本框 → 版面识别/表格识别 → 文本合并 →
    按「题号风格」正则逐框判断哪一行是题目 → 题目之后的行累积成答案 →
    每对 Q&A 从页面渲染图裁出一张截图，同时记录题目在页上的真实坐标。

    输入数据的样子：
        filename / binary —— PDF 文件路径或二进制内容
        from_page / to_page —— 只解析这个页码范围（0 基闭区间）
        zoomin —— 页面渲染图的放大倍数（默认 3，保证裁图清晰）

    输出数据的样子（qai_list，chunk() 里逐对喂给 beAdocPdf）：
        qai_list = [
            ("第一问 什么是RAG？",            # q：问题文本（含题号前缀）
             "RAG是检索增强生成，它……",      # a：该题答案（可能跨多段）
             <PIL.Image>,                   # image：从该题所在页面区域裁出的截图
             [(0, 72.0, 583.0, 100.0, 200.0), ...]),  # poss：题目段落坐标（页号 0 基）
            ...
        ]
        注意页号基准：poss 的页号是 0 基（crop 按 0 基返回），后续 add_positions
        会转成 1 基入库；而 get_tbls_info 的 tbl_pn 是 1 基（拼标签时 +1，crop 解析时再减回 0 基），两者勿混用。
    """

    def __call__(self, filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, zoomin=3, callback=None):
        # —— 第 1 步：OCR 渲染 —— 把 PDF 每页渲染成高清图，并识别出所有文本/版面
        start = timer()
        callback(msg="OCR started")
        self.__images__(filename if not binary else binary, zoomin, from_page, to_page, callback)
        callback(msg="OCR finished ({:.2f}s)".format(timer() - start))
        logging.debug("OCR({}~{}): {:.2f}s".format(from_page, to_page, timer() - start))
        start = timer()

        # —— 第 2 步：版面分析 —— 识别每个文本框的布局类型（标题/正文/图片等）
        self._layouts_rec(zoomin, drop=False)
        callback(0.63, "Layout analysis ({:.2f}s)".format(timer() - start))

        # —— 第 3 步：表格结构识别 —— 找出页面里的表格区域并转成 HTML
        start = timer()
        self._table_transformer_job(zoomin)
        callback(0.65, "Table analysis ({:.2f}s)".format(timer() - start))

        # —— 第 4 步：文本合并 + 表格抽取 —— 把碎片文本按行合并成完整段落，
        #              并从页面中抽取表格/图片（带坐标，供裁图和位置登记用）
        start = timer()
        self._text_merge()
        callback(0.67, "Text merged ({:.2f}s)".format(timer() - start))
        tbls = self._extract_table_figure(True, zoomin, True, True)
        logging.debug("layouts: {}".format(timer() - start))

        # —— 第 5 步：判定 Q&A 结构 —— 取所有文本框文本，按 QUESTION_PATTERN 顺序
        #              找出「第一个命中的题号风格」（先到先得，不是多数投票），
        #              一个都命中不了（不是问答型文档）直接抛异常终止解析
        sections = [b["text"] for b in self.boxes]
        bull_x0_list = []
        q_bull, reg = qbullets_category(sections)
        if q_bull == -1:
            raise ValueError("Unable to recognize Q&A structure.")

        qai_list = []
        last_q, last_a, last_tag = "", "", ""   # 正在累积的 (问题, 答案, 位置标签串)
        last_index = -1                          # 上一题的题号（用于判断递增）
        last_box = {"text": ""}                  # 上一个文本框（位置防误判用）
        last_bull = None                         # 上一个框是否被判为题目

        # 表格/图片按 (页号, 顶部y) 排序，方便与正文行按顺序穿插（图片可能夹在答案中间）
        def sort_key(element):
            tbls_pn = element[1][0][0]
            tbls_top = element[1][0][3]
            return tbls_pn, tbls_top

        tbls.sort(key=sort_key)
        tbl_index = 0
        last_pn, last_bottom = 0, 0
        # 当前正在跟踪的表格/图片的位置信息与文本（@@页\t左\t右\t上\t下## 标签格式）
        tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, tbl_text = 1, 0, 0, 0, 0, "@@0\t0\t0\t0\t0##", ""
        for box in self.boxes:
            # —— 第 6 步：逐框判断「是不是题目」 ——
            #             has_qbullet 综合「题号正则命中 + 题号递增 + 排版位置 +
            #             前段冒号结尾引导语」多重证据（前一段以冒号/问号结尾
            #             通常意味着"题目说明"，该框不算题目）
            section, line_tag = box["text"], self._line_tag(box, zoomin)
            has_bull, index = has_qbullet(reg, box, last_box, last_index, last_bull, bull_x0_list)
            last_box, last_index, last_bull = box, index, has_bull
            line_pn = get_float(line_tag.lstrip("@@").split("\t")[0])
            line_top = get_float(line_tag.rstrip("##").split("\t")[3])
            tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, tbl_text = self.get_tbls_info(tbls, tbl_index)
            if not has_bull:  # 不是题目 → 是答案的一部分
                if not last_q:
                    # 还没有题目在累积（开头的前言部分），跳过；期间路过的表格/图片前进一个
                    if tbl_pn < line_pn or (tbl_pn == line_pn and tbl_top <= line_top):  # image passed
                        tbl_index += 1
                    continue
                else:
                    # 有题目在累积 → 把当前行并进答案；如果中间夹着表格/图片，
                    # 也要把它们按位置顺序拼进答案文本与位置标签
                    sum_tag = line_tag
                    sum_section = section
                    while ((tbl_pn == last_pn and tbl_top >= last_bottom) or (tbl_pn > last_pn)) and (
                        (tbl_pn == line_pn and tbl_top <= line_top) or (tbl_pn < line_pn)
                    ):  # add image at the middle of current answer
                        sum_tag = f"{tbl_tag}{sum_tag}"
                        sum_section = f"{tbl_text}{sum_section}"
                        tbl_index += 1
                        tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, tbl_text = self.get_tbls_info(tbls, tbl_index)
                    last_a = f"{last_a}{sum_section}"
                    last_tag = f"{last_tag}{sum_tag}"
            else:  # 是题目 → 先把上一题收尾，再开新题
                if last_q:
                    # 收尾上一题：先把上一题答案区域内的表格/图片拼进答案，再裁图、登记坐标
                    while ((tbl_pn == last_pn and tbl_top >= last_bottom) or (tbl_pn > last_pn)) and (
                        (tbl_pn == line_pn and tbl_top <= line_top) or (tbl_pn < line_pn)
                    ):  # add image at the end of last answer
                        last_tag = f"{last_tag}{tbl_tag}"
                        last_a = f"{last_a}{tbl_text}"
                        tbl_index += 1
                        tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, tbl_text = self.get_tbls_info(tbls, tbl_index)
                    image, poss = self.crop(last_tag, need_position=True)  # 按位置标签裁出整道题的截图
                    qai_list.append((last_q, last_a, image, poss))
                    last_q, last_a, last_tag = "", "", ""
                # 开新题：题号前缀（"第一问"）当问题，题号后面的文字当答案的开头
                last_q = has_bull.group()
                _, end = has_bull.span()
                last_a = section[end:]
                last_tag = line_tag
            # 记录当前行的页底位置，供上面「图片是否夹在答案里」的穿插判断使用
            last_bottom = float(line_tag.rstrip("##").split("\t")[4])
            last_pn = line_pn
        # 文档末尾可能还有没收尾的最后一题
        if last_q:
            qai_list.append((last_q, last_a, *self.crop(last_tag, need_position=True)))
        return qai_list, tbls

    def get_tbls_info(self, tbls, tbl_index):
        """取出「第 tbl_index 个表格/图片」的位置与文本 —— 表格游标读取器。

        输入数据的样子：
            tbls —— _extract_table_figure 的产物，元素形如
                    ( (image, 文本), [(页号, 左, 右, 上, 下), ...] )
                    文本是字符串=表格HTML，是列表=图片描述行（约定见 tokenize_table）
            tbl_index —— 当前游标（越界时返回占位数据，见下）

        输出：(页号, 左, 右, 上, 下, 位置标签串, 文本)
            位置标签串形如 "@@1\t72.0\t583.0\t100.0\t200.0##"，会被 crop() 解析去裁图。
            注意页号基准：这里返回的 tbl_pn 是 **1 基**（解析器 0 基 → +1），
            拼进 tbl_tag 后 crop/extract_positions 解析时再减 1 回 0 基；
            而 crop 返回给 beAdocPdf 的 poss 页号是 **0 基**，两者勿混用。

        越界占位（tbl_index >= len(tbls)）：返回 tbl_pn=1、坐标全 0、
        标签 "@@0\t0\t0\t0\t0##"（页号 0）、空文本 —— 这是个"无效坐标"哨兵：
        标签里的页号 0 无法命中真实页面，crop 裁不到任何东西；
        同时 tbl_pn=1 无法满足外层 while 的 tbl_pn > last_pn 条件（last_pn >= 1），
        保证「表格插完」后循环自然终止，不会越界空转。
        """
        if tbl_index >= len(tbls):
            # 越界 → 返回占位：页号 1、坐标全 0、标签全 0、无文本（等效"没有表格"）
            return 1, 0, 0, 0, 0, "@@0\t0\t0\t0\t0##", ""
        tbl_pn = tbls[tbl_index][1][0][0] + 1   # 表格所在页号（解析器 0 基 → 这里转 1 基）
        tbl_left = tbls[tbl_index][1][0][1]
        tbl_right = tbls[tbl_index][1][0][2]
        tbl_top = tbls[tbl_index][1][0][3]
        tbl_bottom = tbls[tbl_index][1][0][4]
        tbl_tag = "@@{}\t{:.1f}\t{:.1f}\t{:.1f}\t{:.1f}##".format(tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom)
        _tbl_text = "".join(tbls[tbl_index][0][1])
        return tbl_pn, tbl_left, tbl_right, tbl_top, tbl_bottom, tbl_tag, _tbl_text


class Docx(DocxParser):
    """DOCX 版 Q&A 解析器 —— 把「Word 问答文档」读成 (问题, 答案, 截图) 列表。

    处理链（一句话）：逐段落读取，用「Heading 标题样式」识别题目（多级题号用栈维护）→
    非标题段落累积成答案 → 段落里的图片拼到答案尾部 → 文档原生表格转成 HTML 交给下游。

    输入数据的样子：
        filename / binary —— docx 文件路径或二进制内容
        from_page / to_page —— 只解析这个页码范围（与 PDF 分支保持一致的参数形态）

    输出数据的样子（qai_list，chunk() 里逐对喂给 beAdocDocx）：
        qai_list = [
            ("1.1",                          # q：完整题号（多级用 \n 连接）
             "什么是RAG？\nRAG是检索增强生成……",  # a：该题全部正文（含问题正文）
             <PIL.Image>),                   # image：该题段落里的图片（可能 None）
            ...
        ]
        以及 tbls：文档里的原生表格，转成 ((None, "<table>…</table>"), "") 形式
    """

    def __init__(self):
        pass

    def __call__(self, filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, callback=None):
        self.doc = Document(filename) if not binary else Document(BytesIO(binary))
        pn = 0  # 当前页码（遇到分页符 +1，用来配合 from_page/to_page 过滤）
        last_answer, last_image = "", None  # 正在累积的 (答案文本, 图片)
        question_stack, level_stack = [], []  # 多级题号栈 + 每级的样式层级
        qai_list = []
        for p in self.doc.paragraphs:
            if pn > to_page:  # 超出解析范围，直接终止
                break
            question_level, p_text = 0, ""
            # 只在目标页码区间内、且段落有文字时才判定「是不是题目」；
            # docx_question_level(p) 默认 bull=-1，只有 Heading 样式才返回层级，
            # 其余一律 0（当正文处理）。
            # 注意：pn == to_page 这一页不判题，但正文仍会走 else 累积进答案
            if from_page <= pn < to_page and p.text.strip():
                question_level, p_text = docx_question_level(p)
            if not question_level or question_level > 6:  # 正文（超过 6 级也不当题）
                last_answer = f"{last_answer}\n{p_text}"  # 累积进答案
                current_image = self.get_picture(self.doc, p)  # 取该段落的图片
                last_image = concat_img(last_image, current_image)  # 与之前图片拼接（None/同图会短路返回原值）
            else:  # 是题目 → 先收尾上一题，再入栈开新题
                if last_answer or last_image:
                    sum_question = "\n".join(question_stack)  # 用栈拼出完整题号
                    if sum_question:
                        qai_list.append((sum_question, last_answer, last_image))
                    last_answer, last_image = "", None  # 清空，开始累积下一题

                i = question_level
                # 出栈：新题层级 <= 栈顶层级时，说明同级的另一题出现，旧题结束
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(p_text)  # 新题号进栈
                level_stack.append(question_level)  # 记下这级的样式层级
            # 分页符检测：run 里出现 lastRenderedPageBreak 或 type="page" 的换页符 → 页码 +1
            for run in p.runs:
                if "lastRenderedPageBreak" in run._element.xml:
                    pn += 1
                    continue
                if "w:br" in run._element.xml and 'type="page"' in run._element.xml:
                    pn += 1
        if last_answer:  # 文档末尾可能还有没收尾的最后一题
            sum_question = "\n".join(question_stack)
            if sum_question:
                qai_list.append((sum_question, last_answer, last_image))

        # 文档里的原生表格 → 转成 HTML（相邻同文单元格合并成 colspan）
        tbls = []
        for tb in self.doc.tables:
            html = "<table>"
            for r in tb.rows:
                html += "<tr>"
                i = 0
                while i < len(r.cells):
                    span = 1
                    c = r.cells[i]
                    # 向后扫「文本相同」的连续单元格，把它们合并成一个 colspan
                    for j in range(i + 1, len(r.cells)):
                        if c.text == r.cells[j].text:
                            span += 1
                            i = j
                        else:
                            break
                    i += 1
                    html += f"<td>{c.text}</td>" if span == 1 else f"<td colspan='{span}'>{c.text}</td>"
                html += "</tr>"
            html += "</table>"
            tbls.append(((None, html), ""))  # 按 tokenize_table 约定：字符串=表格（整表一个切片）
        return qai_list, tbls


def rmPrefix(txt):
    """剥掉 Q&A 文本开头的「问题/答案/Q/A」等前缀标签 —— 前缀清理工。

    输入数据的样子：
        txt = "问题：什么是RAG？" / "Answer: RAG是检索增强生成…" / "1. 什么是RAG？"

    输出：
        去掉前缀后的干净文本，如 "什么是RAG？"。
        匹配规则：开头允许出现 (问题|答案|回答|user|assistant|Q|A|Question|Answer|问|答)
        之一（大小写不敏感），后面跟 \t 或 : 或 ： 或空格，把这一整段剥掉；
        其余文本原样保留。没匹配到前缀就原样返回。
    """
    return re.sub(r"^(问题|答案|回答|user|assistant|Q|A|Question|Answer|问|答)[\t:： ]+", "", txt.strip(), flags=re.IGNORECASE)


def beAdocPdf(d, q, a, eng, image, poss):
    """把 PDF 解析出的 (问题, 答案, 截图, 坐标) 组装成可入库的切片字典 —— 收口成文档。

    输入数据的样子：
        d = {"docnm_kwd": "试卷.pdf", "title_tks": "试卷"}
        q = "第一问 什么是RAG？"；a = "RAG是检索增强生成……"
        image = <PIL.Image>（该题截图，可能是 None）
        poss = [(0, 72.0, 583.0, 100.0, 200.0), ...]
              （题目段落坐标：crop 按 ZM 换算、中间段宽度修正后的值，
                页号 0 基 —— add_positions 会转成 1 基入库）

    输出（在原 d 上原地补字段后返回 d）：
        d = {
            "docnm_kwd": "试卷.pdf", "title_tks": "试卷",
            "content_with_weight": "问题：第一问 什么是RAG？\t回答：RAG是检索增强生成……",
            "content_ltks": [分词列表], "content_sm_ltks": [细分词列表],
            "image": <PIL.Image>, "doc_type_kwd": "image",   # image 非空才加
            "page_num_int": [1, ...], "position_int": [...], "top_int": [...],  # 来自 add_positions
        }

    干了 4 件事：
        ① 拼 content_with_weight：问题/答案各加前缀（英文 "Question:/Answer:"，
           中文 "问题：/回答："），用 \t 连接成一行
        ② 粗分词 → content_ltks（问题文本），再细分词 → content_sm_ltks
        ③ image 非空时挂上 image 和 doc_type_kwd="image"
        ④ 调 add_positions 把 poss 坐标写进 page_num_int / position_int / top_int
    """
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join([qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    if image:
        d["image"] = image
        d["doc_type_kwd"] = "image"
    add_positions(d, poss)
    return d


def beAdocDocx(d, q, a, eng, image, row_num=-1):
    """把 DOCX 解析出的 (问题, 答案, 截图) 组装成可入库的切片字典 —— 收口成文档。

    与 beAdocPdf 几乎一样，区别只有两点：
        ① 没有 poss 坐标，靠 row_num（题目出现的行号）写 top_int；
        ② row_num < 0 时连 top_int 也不写（pdf/md 等分支会走别的方式登记位置）。

    输入数据的样子：
        d = {"docnm_kwd": "问答.docx", "title_tks": "问答"}
        q = "1.1"；a = "什么是RAG？\nRAG是检索增强生成……"；image = None；row_num = 3

    输出：同 beAdocPdf，另带 "top_int": [3]。
    """
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join([qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    if image:
        d["image"] = image
        d["doc_type_kwd"] = "image"
    if row_num >= 0:
        d["top_int"] = [row_num]
    return d


def beAdoc(d, q, a, eng, row_num=-1):
    """把纯文本解析出的 (问题, 答案) 组装成可入库的切片字典 —— 收口成文档（无图版）。

    用于 Excel / txt / csv / Markdown 分支：这些格式没有截图和物理坐标，
    只有 row_num 行号可以登记位置。

    输入数据的样子：
        d = {"docnm_kwd": "问答.xlsx", "title_tks": "问答"}
        q = "什么是RAG？"；a = "RAG是检索增强生成……"；eng = False；row_num = 0

    输出：同 beAdocDocx（content_with_weight / content_ltks / content_sm_ltks，
        可选 top_int=[row_num]）。
    """
    qprefix = "Question: " if eng else "问题："
    aprefix = "Answer: " if eng else "回答："
    d["content_with_weight"] = "\t".join([qprefix + rmPrefix(q), aprefix + rmPrefix(a)])
    d["content_ltks"] = rag_tokenizer.tokenize(q)
    d["content_sm_ltks"] = rag_tokenizer.fine_grained_tokenize(d["content_ltks"])
    if row_num >= 0:
        d["top_int"] = [row_num]
    return d


def mdQuestionLevel(s):
    """判断 Markdown 一行是不是标题、是第几级 —— Markdown 层级识别员。

    输入数据的样子：
        s = "### 什么是RAG？"  /  "1. 什么是RAG？"（正文）  /  "   普通正文"

    输出：
        (层级数, 去掉 # 和行首空格的文本)
            "### 什么是RAG？" → (3, "什么是RAG？")
            "普通正文"       → (0, "普通正文")（没有 # 就不算标题）
        层级 1~6 对应 Markdown 的 H1~H6，0 表示正文。
    """
    match = re.match(r"#*", s)
    return (len(match.group(0)), s.lstrip("#").lstrip()) if match else (0, s)


def chunk(filename, binary=None, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, lang="Chinese", callback=None, **kwargs):
    """Q&A 文档解析总入口 —— 按扩展名分派到六种解析器，产出可入库的切片列表。

    约定：
        Excel 格式：两列、无表头，第 1 列问题、第 2 列答案，允许多 sheet；
        csv/txt 格式：UTF-8 编码，用 TAB（csv 也接受逗号）分隔问题和答案；
        pdf 格式：试卷/问答册式排版，自动识别题号风格；
        markdown 格式：按 # 标题层级组织 Q&A（代码块内的 # 不算标题）；
        docx 格式：按 Heading 标题样式组织 Q&A。
    所有格式中畸形行都会被忽略；每一对 Q&A 作为一个切片（chunk）。

    输入数据的样子：
        filename = "问答.xlsx"（或 .txt/.csv/.pdf/.md/.docx）
        binary = 文件二进制内容（可选，给了就绕过本地路径）
        from_page / to_page —— 仅 PDF 分支使用，控制页码范围
        lang = "Chinese"（或 "English"）—— 决定分词语言与 Q/A 前缀

    输出数据的样子（每个元素可直接入 ES 的切片字典）：
        res = [
            {
                "docnm_kwd": "问答.xlsx",
                "title_tks": "问答",
                "content_with_weight": "问题：什么是RAG？\t回答：RAG是检索增强生成……",
                "content_ltks": [...分词...],
                "content_sm_ltks": [...细分词...],
                "top_int": [0],        # 行号/位置（Excel/txt/csv/docx 有，md 无）
            },
            ...
        ]
        PDF 分支还会带上 "image"、"doc_type_kwd": "image"、page_num_int/position_int。
        （Excel 的 q 若带"问题："前缀会在 beAdoc 里被 rmPrefix 剥掉，示例按剥后结果展示）

    干了 7 件事：
        ① 按 lang 判断是否英文，并设置分词器的语言
        ② 初始化基础 doc 字典：docnm_kwd=文件名、title_tks=去扩展名后的文件名分词
        ③ 六种格式各自解析（Excel / txt / csv / pdf / md / docx），
           每种都产出 (问题, 答案) 对并用对应的 beAdoc* 收口成切片字典
        ④ txt 分支：统计逗号/TAB 出现次数选分隔符，答案多行时用 \n 累积
        ⑤ csv 分支：用 csv.reader 处理带引号的跨行字段（reader.line_num 追踪物理行）
        ⑥ 全程回调进度（0.1 开始 → 0.6 解析完成 → 1.0 完成）
        ⑦ 不认识的扩展名抛 NotImplementedError
    """
    eng = lang.lower() == "english"
    rag_tokenizer.tokenizer.set_language(lang)  # 设置分词语言（中文/英文）
    res = []
    doc = {"docnm_kwd": filename, "title_tks": rag_tokenizer.tokenize(re.sub(r"\.[a-zA-Z]+$", "", filename))}
    if re.search(r"\.xlsx?$", filename, re.IGNORECASE):
        # —— Excel 分支 —— 每行前两个非空单元格 = (问题, 答案)，beAdoc 收口（带行号）
        callback(0.1, "Start to parse.")
        excel_parser = Excel()
        for ii, (q, a) in enumerate(excel_parser(filename, binary, callback)):
            res.append(beAdoc(deepcopy(doc), q, a, eng, ii))
        return res

    elif re.search(r"\.(txt)$", filename, re.IGNORECASE):
        # —— txt 分支 —— 用「逗号 vs TAB 哪个命中更多」猜分隔符；
        #              一行正好拆成 2 段才算一对 Q&A，否则若已有问题就并入答案（多行答案）
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        comma, tab = 0, 0
        for line in lines:
            if len(line.split(",")) == 2:
                comma += 1
            if len(line.split("\t")) == 2:
                tab += 1
        delimiter = "\t" if tab >= comma else ","

        fails = []
        question, answer = "", ""
        i = 0
        while i < len(lines):
            arr = lines[i].split(delimiter)
            if len(arr) != 2:
                if question:
                    answer += "\n" + lines[i]  # 拆不成一对 → 当作答案的续行
                else:
                    fails.append(str(i + 1))  # 开头就是畸形行 → 记失败
            elif len(arr) == 2:
                if question and answer:
                    res.append(beAdoc(deepcopy(doc), question, answer, eng, i))
                question, answer = arr  # 新的一对 Q&A
            i += 1
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question:  # 收尾最后一对
            res.append(beAdoc(deepcopy(doc), question, answer, eng, len(lines)))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        return res

    elif re.search(r"\.(csv)$", filename, re.IGNORECASE):
        # —— csv 分支 —— 与 txt 类似，但用 csv.reader 解析：
        #              带引号的字段可以跨多行，reader.line_num 记录真实物理行
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        delimiter = "\t" if any("\t" in line for line in lines) else ","

        fails = []
        question, answer = "", ""
        res = []
        reader = csv.reader((line + "\n" for line in lines), delimiter=delimiter)
        prev_line_num = 0

        # 一个「逻辑行」可能吃掉多行物理行（引号内换行），
        # raw 取回这段原文，多行答案用它拼回去
        for i, row in enumerate(reader):
            raw = "\n".join(lines[prev_line_num : reader.line_num])
            prev_line_num = reader.line_num
            if len(row) != 2:
                if question:
                    answer += "\n" + raw  # 畸形/跨行 → 并入答案
                else:
                    fails.append(str(i + 1))
            elif len(row) == 2:
                if question and answer:
                    res.append(beAdoc(deepcopy(doc), question, answer, eng, i))
                question, answer = row
            if len(res) % 999 == 0:
                callback(len(res) * 0.6 / len(lines), ("Extract Q&A: {}".format(len(res)) + (f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))

        if question:
            res.append(beAdoc(deepcopy(doc), question, answer, eng, len(lines)))

        callback(0.6, ("Extract Q&A: {}".format(len(res)) + (f"{len(fails)} failure, line: %s..." % (",".join(fails[:3])) if fails else "")))
        return res

    elif re.search(r"\.pdf$", filename, re.IGNORECASE):
        # —— PDF 分支 —— Pdf 解析器产出 (q, a, 截图, 坐标)，beAdocPdf 收口
        callback(0.1, "Start to parse.")
        pdf_parser = Pdf()
        qai_list, tbls = pdf_parser(filename if not binary else binary, from_page=from_page, to_page=to_page, callback=callback)
        for q, a, image, poss in qai_list:
            res.append(beAdocPdf(deepcopy(doc), q, a, eng, image, poss))
        return res

    elif re.search(r"\.(md|markdown|mdx)$", filename, re.IGNORECASE):
        # —— Markdown 分支 —— # 标题 = 题目（多级栈），正文 = 答案；
        #              ``` 代码块内的行不当题目；答案用 markdown() 渲染成 HTML
        callback(0.1, "Start to parse.")
        txt = get_text(filename, binary)
        lines = txt.split("\n")
        _last_question, last_answer = "", ""
        question_stack, level_stack = [], []
        code_block = False
        for index, line in enumerate(lines):
            if line.strip().startswith("```"):
                code_block = not code_block  # 代码块开关（块内 # 不算标题）
            question_level, question = 0, ""
            if not code_block:
                question_level, question = mdQuestionLevel(line)

            if not question_level or question_level > 6:  # 正文 → 累积答案
                last_answer = f"{last_answer}\n{line}"
            else:  # 标题 → 收尾上一题，再入栈开新题
                if last_answer.strip():
                    sum_question = "\n".join(question_stack)
                    if sum_question:
                        # 答案先渲染成 HTML（支持表格扩展），再收口成切片
                        res.append(beAdoc(deepcopy(doc), sum_question, markdown(last_answer, extensions=["markdown.extensions.tables"]), eng, index))
                    last_answer = ""

                i = question_level
                # 出栈：新标题层级 <= 栈顶 → 同级/更浅标题出现，旧题结束
                while question_stack and i <= level_stack[-1]:
                    question_stack.pop()
                    level_stack.pop()
                question_stack.append(question)
                level_stack.append(question_level)
        if last_answer.strip():  # 收尾最后一题
            sum_question = "\n".join(question_stack)
            if sum_question:
                res.append(beAdoc(deepcopy(doc), sum_question, markdown(last_answer, extensions=["markdown.extensions.tables"]), eng, index))
        return res

    elif re.search(r"\.docx$", filename, re.IGNORECASE):
        # —— DOCX 分支 —— Docx 解析器产出 (q, a, 截图)；
        #              文档表格 tbls 先经 tokenize_table 收口成切片（整表一个切片），
        #              再逐对 Q&A 用 beAdocDocx 收口
        docx_parser = Docx()
        qai_list, tbls = docx_parser(filename, binary, from_page=0, to_page=MAXIMUM_PAGE_NUMBER, callback=callback)
        res = tokenize_table(tbls, doc, eng, language=lang)
        for i, (q, a, image) in enumerate(qai_list):
            res.append(beAdocDocx(deepcopy(doc), q, a, eng, image, i))
        return res

    raise NotImplementedError("Excel, csv(txt), pdf, markdown and docx format files are supported.")


if __name__ == "__main__":
    import sys

    def dummy(prog=None, msg=""):
        pass

    chunk(sys.argv[1], from_page=0, to_page=10, callback=dummy)
